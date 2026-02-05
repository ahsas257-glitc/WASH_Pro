from __future__ import annotations

import os
import tempfile
import hashlib
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple, Callable
from urllib.parse import urlparse

import streamlit as st
from PIL import Image

from src.Tools.utils.types import Tool6Context
from design.components.base_tool_ui import card_open, card_close, status_card


# =============================================================================
# Session keys (Tool 6 naming)
# =============================================================================
SS_OBS = "tool6_obs_components"  # list of components (UI + report-ready)
SS_PHOTO_BYTES = "photo_bytes"   # shared cache, used in final report
SS_AUDIO_BYTES = "audio_bytes"   # optional local cache

# Preview cache (performance)
SS_OBS_PREVIEW_DOCX_HASH = "tool6_obs_preview_docx_hash"
SS_OBS_PREVIEW_DOCX_BYTES = "tool6_obs_preview_docx_bytes"
SS_OBS_PREVIEW_PNG_HASH = "tool6_obs_preview_png_hash"
SS_OBS_PREVIEW_PNG_BYTES = "tool6_obs_preview_png_bytes"


# =============================================================================
# Titles
# =============================================================================
DEFAULT_OBSERVATION_TITLES: List[str] = [
    "Construction of bore well and well protection structure:",
    "Supply and installation of the solar system:",
    "Construction of 60 m3 reservoir:",
    "Construction of 5 m3 reservoir for School:",
    "Construction of boundary wall:",
    "Construction of guard room and latrine:",
    "Construction of stand taps:",
]
DEFAULT_OBSERVATION_TITLES = list(dict.fromkeys(DEFAULT_OBSERVATION_TITLES))


# =============================================================================
# Layout constants
# =============================================================================
PHOTO_W_IN = 3.17
PHOTO_H_IN = 2.38
PHOTO_ASPECT = PHOTO_W_IN / PHOTO_H_IN  # ~1.3328


# =============================================================================
# Small helpers
# =============================================================================
def _s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def _k(*parts: Any) -> str:
    # short, stable
    raw = ".".join(str(p) for p in parts)
    h = hashlib.md5(raw.encode("utf-8")).hexdigest()[:10]
    return f"t6.s3.{h}"


def _sha1(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()


def _ensure_state() -> None:
    ss = st.session_state

    ss.setdefault(SS_OBS, [])
    if not isinstance(ss[SS_OBS], list):
        ss[SS_OBS] = []

    ss.setdefault(SS_PHOTO_BYTES, {})
    if not isinstance(ss[SS_PHOTO_BYTES], dict):
        ss[SS_PHOTO_BYTES] = {}

    ss.setdefault(SS_AUDIO_BYTES, {})
    if not isinstance(ss[SS_AUDIO_BYTES], dict):
        ss[SS_AUDIO_BYTES] = {}


def _ensure_component_schema(c: Dict[str, Any]) -> Dict[str, Any]:
    c.setdefault("comp_id", "")
    c.setdefault("title", "")
    c.setdefault("observations", [])
    c.setdefault("observations_valid", [])
    return c


def _ensure_obs_schema(it: Dict[str, Any]) -> Dict[str, Any]:
    it.setdefault("title_mode", "Select")
    it.setdefault("title_selected", "")
    it.setdefault("title_custom", "")
    it.setdefault("text", "")
    it.setdefault("audio_url", "")
    it.setdefault("photos", [])
    return it


def _obs_title_raw(it: Dict[str, Any]) -> str:
    it = _ensure_obs_schema(it)
    if it.get("title_mode") == "Custom":
        return _s(it.get("title_custom"))
    return _s(it.get("title_selected"))


def _numbered_title(section_no: str, global_idx_1based: int, raw_title: str) -> str:
    t = _s(raw_title)
    if not t:
        return ""
    return f"{section_no}.{global_idx_1based}. {t}"


def _normalize_photos(selected_urls: List[str], old_photos: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    old_map = {
        _s(p.get("url")): _s(p.get("note"))
        for p in (old_photos or [])
        if isinstance(p, dict) and _s(p.get("url"))
    }
    return [{"url": u, "note": old_map.get(u, "")} for u in selected_urls]


# =============================================================================
# Mutations
# =============================================================================
def _add_component() -> None:
    st.session_state[SS_OBS].append(
        _ensure_component_schema(
            {
                "comp_id": "",
                "title": "",
                "observations": [_ensure_obs_schema({})],
                "observations_valid": [],
            }
        )
    )


def _remove_component(idx: int) -> None:
    comps = st.session_state.get(SS_OBS, [])
    if isinstance(comps, list) and 0 <= idx < len(comps):
        comps.pop(idx)
        st.session_state[SS_OBS] = comps


def _add_observation(comp_idx: int) -> None:
    comps = st.session_state.get(SS_OBS, [])
    if not (isinstance(comps, list) and 0 <= comp_idx < len(comps)):
        return
    comp = _ensure_component_schema(comps[comp_idx])
    obs = comp.get("observations") or []
    obs.append(_ensure_obs_schema({}))
    comp["observations"] = obs
    comps[comp_idx] = comp
    st.session_state[SS_OBS] = comps


def _remove_observation(comp_idx: int, obs_idx: int) -> None:
    comps = st.session_state.get(SS_OBS, [])
    if not (isinstance(comps, list) and 0 <= comp_idx < len(comps)):
        return
    comp = _ensure_component_schema(comps[comp_idx])
    obs = comp.get("observations") or []
    if 0 <= obs_idx < len(obs):
        obs.pop(obs_idx)
    if not obs:
        obs = [_ensure_obs_schema({})]
    comp["observations"] = obs
    comps[comp_idx] = comp
    st.session_state[SS_OBS] = comps


# =============================================================================
# Media caching
# =============================================================================
def _fetch_and_cache_image(
    url: str,
    *,
    fetch_image: Callable[[str], Tuple[bool, Optional[bytes], str]],
) -> None:
    url = _s(url)
    if not url:
        return

    photo_bytes: Dict[str, bytes] = st.session_state.get(SS_PHOTO_BYTES, {})
    if url in photo_bytes and photo_bytes[url]:
        return

    ok, b, _msg = fetch_image(url)
    if ok and b:
        photo_bytes[url] = b
        st.session_state[SS_PHOTO_BYTES] = photo_bytes


def _fetch_and_cache_audio(
    url: str,
    *,
    fetch_audio: Optional[Callable[[str], Tuple[bool, Optional[bytes], str, str]]],
) -> Tuple[Optional[bytes], str]:
    url = _s(url)
    if not url or fetch_audio is None:
        return None, ""

    audio_bytes: Dict[str, bytes] = st.session_state.get(SS_AUDIO_BYTES, {})
    if url in audio_bytes and audio_bytes[url]:
        return audio_bytes[url], "audio/aac"

    ok, b, _msg, mime = fetch_audio(url)
    if ok and b:
        audio_bytes[url] = b
        st.session_state[SS_AUDIO_BYTES] = audio_bytes
        return b, (mime or "audio/aac").split(";")[0]

    return None, ""


# =============================================================================
# Build valid observations with GLOBAL numbering
# =============================================================================
def _build_valid_observations_global(
    section_no: str,
    observations: List[Dict[str, Any]],
    *,
    start_index_1based: int,
    photo_bytes_cache: Dict[str, bytes],
) -> Tuple[List[Dict[str, Any]], int]:
    valid: List[Dict[str, Any]] = []
    global_idx = int(start_index_1based)

    for it in (observations or []):
        if not isinstance(it, dict):
            continue
        it = _ensure_obs_schema(it)

        title_raw = _obs_title_raw(it)
        if not _s(title_raw):
            continue

        title_num = _numbered_title(section_no, global_idx, title_raw)
        if not title_num:
            continue

        photos_fixed: List[Dict[str, Any]] = []
        for p in (it.get("photos") or []):
            if isinstance(p, dict) and _s(p.get("url")):
                u = _s(p.get("url"))
                photos_fixed.append(
                    {
                        "url": u,
                        "note": _s(p.get("note")),
                        "bytes": photo_bytes_cache.get(u),
                    }
                )

        valid.append(
            {
                "title": title_num,
                "text": _s(it.get("text")),
                "audio_url": _s(it.get("audio_url")),
                "photos": photos_fixed,
            }
        )

        global_idx += 1

    return valid, global_idx


# =============================================================================
# Preview helpers (LAZY IMPORTS for speed)
# =============================================================================
def _crop_to_aspect(img: Image.Image, target_aspect: float) -> Image.Image:
    w, h = img.size
    if w <= 0 or h <= 0:
        return img
    current = w / h
    if abs(current - target_aspect) < 1e-3:
        return img
    if current > target_aspect:
        new_w = int(h * target_aspect)
        left = max(0, (w - new_w) // 2)
        return img.crop((left, 0, left + new_w, h))
    new_h = int(w / target_aspect)
    top = max(0, (h - new_h) // 2)
    return img.crop((0, top, w, top + new_h))


def _to_clean_png_fit_box(img_bytes: bytes, *, target_aspect: float) -> Optional[bytes]:
    try:
        img = Image.open(BytesIO(img_bytes)).convert("RGB")
        img = _crop_to_aspect(img, target_aspect)
        out = BytesIO()
        img.save(out, format="PNG", optimize=True)
        return out.getvalue()
    except Exception:
        return None


def _docx_first_page_to_png(docx_bytes: bytes) -> Optional[bytes]:
    """
    Exact Word-like PNG needs Windows + MS Word + pywin32 + pymupdf.
    Lazy imports to keep initial load fast.
    """
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
        import fitz  # type: ignore
    except Exception:
        return None

    with tempfile.TemporaryDirectory() as td:
        docx_path = os.path.join(td, "preview.docx")
        pdf_path = os.path.join(td, "preview.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            doc = word.Documents.Open(
                docx_path,
                ReadOnly=True,
                AddToRecentFiles=False,
                ConfirmConversions=False,
                NoEncodingDialog=True,
            )

            try:
                doc.SaveAs2(pdf_path, FileFormat=17)
            except Exception:
                doc.SaveAs(pdf_path, FileFormat=17)

            doc.Close(False)
            doc = None
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            pythoncom.CoUninitialize()

        try:
            pdf = fitz.open(pdf_path)
            if pdf.page_count < 1:
                return None
            page = pdf.load_page(0)
            pix = page.get_pixmap(dpi=170)
            out = pix.tobytes("png")
            pdf.close()
            return out
        except Exception:
            return None


def _build_observations_preview_docx_bytes(
    *,
    ctx: Tool6Context,
    comps: List[Dict[str, Any]],
    photo_bytes_cache: Dict[str, bytes],
) -> bytes:
    # Lazy import: python-docx only when preview is enabled
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _set_para_compact(p) -> None:
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0

    def _set_cell_margins(cell, top=0, start=80, bottom=0, end=80):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is None:
            tcMar = OxmlElement("w:tcMar")
            tcPr.append(tcMar)

        for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tcMar.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                tcMar.append(node)
            node.set(qn("w:w"), str(int(val)))
            node.set(qn("w:type"), "dxa")

    def _remove_table_borders(table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = tblBorders.find(qn(f"w:{edge}"))
            if elem is None:
                elem = OxmlElement(f"w:{edge}")
                tblBorders.append(elem)
            elem.set(qn("w:val"), "nil")

    doc = Document()

    p = doc.add_paragraph("5. Project Component-Wise Observations (Preview)")
    if p.runs:
        p.runs[0].bold = True
    _set_para_compact(p)

    doc.add_paragraph("")

    for comp in (comps or [])[:4]:
        if not isinstance(comp, dict):
            continue

        comp_title = _s(comp.get("title")) or "Component"
        comp_id = _s(comp.get("comp_id"))
        hdr = f"{comp_id} — {comp_title}".strip(" —")

        ph = doc.add_paragraph(hdr)
        try:
            ph.style = "Heading 2"
        except Exception:
            if ph.runs:
                ph.runs[0].bold = True
        _set_para_compact(ph)

        obs_valid = comp.get("observations_valid") or []
        for it in obs_valid[:6]:
            if not isinstance(it, dict):
                continue

            t = _s(it.get("title"))
            txt = _s(it.get("text"))
            photos = it.get("photos") or []

            th = doc.add_paragraph(t)
            try:
                th.style = "Heading 3"
            except Exception:
                pass
            _set_para_compact(th)

            tbl = doc.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            _remove_table_borders(tbl)

            left = tbl.rows[0].cells[0]
            right = tbl.rows[0].cells[1]
            left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            _set_cell_margins(left, start=60, end=160)
            _set_cell_margins(right, start=160, end=60)

            p1 = left.paragraphs[0]
            p1.text = txt
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_para_compact(p1)

            right.paragraphs[0].text = ""
            _set_para_compact(right.paragraphs[0])

            ph_list = [p for p in (photos or []) if isinstance(p, dict) and _s(p.get("url"))]
            if ph_list:
                first = ph_list[0]
                u = _s(first.get("url"))
                b = first.get("bytes") or photo_bytes_cache.get(u)
                if b:
                    clean = _to_clean_png_fit_box(b, target_aspect=PHOTO_ASPECT)
                    if clean:
                        pic_p = right.add_paragraph()
                        pic_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _set_para_compact(pic_p)
                        pic_p.add_run().add_picture(BytesIO(clean), width=Inches(PHOTO_W_IN), height=Inches(PHOTO_H_IN))

            doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _get_or_build_obs_preview_docx(
    ctx: Tool6Context,
    comps: List[Dict[str, Any]],
    photo_bytes_cache: Dict[str, bytes],
) -> Optional[bytes]:
    selected_urls: List[str] = []
    minimal = []

    for comp in comps or []:
        if not isinstance(comp, dict):
            continue
        ov = comp.get("observations_valid") or []
        minimal.append(
            {
                "comp_id": _s(comp.get("comp_id")),
                "title": _s(comp.get("title")),
                "ov": [
                    {
                        "title": _s(o.get("title")),
                        "text": _s(o.get("text")),
                        "photos": [(_s(p.get("url")), _s(p.get("note"))) for p in (o.get("photos") or [])],
                    }
                    for o in ov
                    if isinstance(o, dict)
                ],
            }
        )
        for o in ov:
            for p in (o.get("photos") or []):
                u = _s(p.get("url"))
                if u:
                    selected_urls.append(u)

    ph = []
    for u in sorted(set(selected_urls)):
        b = photo_bytes_cache.get(u) or b""
        ph.append((u, _sha1(b) if b else ""))

    fp = repr((minimal, ph)).encode("utf-8")
    h = _sha1(fp)

    if st.session_state.get(SS_OBS_PREVIEW_DOCX_HASH) == h and st.session_state.get(SS_OBS_PREVIEW_DOCX_BYTES):
        return st.session_state[SS_OBS_PREVIEW_DOCX_BYTES]

    try:
        docx_bytes = _build_observations_preview_docx_bytes(ctx=ctx, comps=comps, photo_bytes_cache=photo_bytes_cache)
    except Exception:
        return None

    st.session_state[SS_OBS_PREVIEW_DOCX_HASH] = h
    st.session_state[SS_OBS_PREVIEW_DOCX_BYTES] = docx_bytes
    st.session_state[SS_OBS_PREVIEW_PNG_HASH] = None
    st.session_state[SS_OBS_PREVIEW_PNG_BYTES] = None
    return docx_bytes


def _get_or_build_obs_preview_png(docx_bytes: bytes) -> Optional[bytes]:
    if not docx_bytes:
        return None
    h = _sha1(docx_bytes)

    if st.session_state.get(SS_OBS_PREVIEW_PNG_HASH) == h and st.session_state.get(SS_OBS_PREVIEW_PNG_BYTES):
        return st.session_state[SS_OBS_PREVIEW_PNG_BYTES]

    png = _docx_first_page_to_png(docx_bytes)
    if png:
        st.session_state[SS_OBS_PREVIEW_PNG_HASH] = h
        st.session_state[SS_OBS_PREVIEW_PNG_BYTES] = png
    return png


# =============================================================================
# MAIN RENDER
# =============================================================================
def render_step(
    ctx: Tool6Context,
    *,
    fetch_image: Callable[[str], Tuple[bool, Optional[bytes], str]],
    fetch_audio: Optional[Callable[[str], Tuple[bool, Optional[bytes], str, str]]] = None,
) -> bool:
    _ensure_state()

    st.subheader("Step 3 — Observations (Audio + Photos + Notes)")

    urls = getattr(ctx, "all_photo_urls", []) or []
    labels = getattr(ctx, "photo_label_by_url", {}) or {}
    audios = getattr(ctx, "audios", []) or []

    audio_url_list: List[str] = []
    audio_label_by_url: Dict[str, str] = {}
    for i, a in enumerate(audios, start=1):
        if isinstance(a, dict):
            u = _s(a.get("url"))
            f = _s(a.get("field")) or "Audio"
            if u:
                audio_url_list.append(u)
                audio_label_by_url[u] = f"{i:02d}. {f}"

    with st.container(border=True):
        card_open(
            "Observations",
            subtitle=(
                "Select/write a title, choose photos, and write a note for each photo. "
                "Only observations with a title will be included in the report."
            ),
            variant="lg-variant-green",
        )

        comps: List[Dict[str, Any]] = st.session_state[SS_OBS]

        a, b, c, d = st.columns([1, 1, 1, 1], gap="small")
        with a:
            st.button("Add component", use_container_width=True, key=_k("add_comp"), on_click=_add_component)
        with b:
            show_previews = st.toggle("Show photo previews", value=False, key=_k("show_prev"))
        with c:
            preview_on = st.toggle("Show Word-like preview", value=False, key=_k("show_word_preview"))
        with d:
            if st.button("Clear all", use_container_width=True, key=_k("clear_all")):
                st.session_state[SS_OBS] = []
                comps = st.session_state[SS_OBS]

        if not comps:
            status_card("No components yet", "Click **Add component** to start.", level="warning")
            card_close()
            return False

        st.divider()

        SECTION_NO = "5"
        global_obs_idx = 1  # global across all components

        for ci in range(len(comps)):
            comp = _ensure_component_schema(comps[ci])

            comp_title = _s(comp.get("title")) or f"Component {ci + 1}"
            comp_id = _s(comp.get("comp_id"))
            exp_title = f"{comp_id} — {comp_title}".strip(" —")

            with st.expander(exp_title, expanded=(ci == 0)):
                top = st.columns([1, 1, 2], gap="small")
                with top[0]:
                    st.button(
                        "Remove component",
                        use_container_width=True,
                        key=_k("rm_comp", ci),
                        on_click=_remove_component,
                        args=(ci,),
                    )
                with top[1]:
                    st.button(
                        "Add observation",
                        use_container_width=True,
                        key=_k("add_obs", ci),
                        on_click=_add_observation,
                        args=(ci,),
                    )
                with top[2]:
                    st.caption("You can add multiple observations under this component.")

                c1, c2 = st.columns([1, 1], gap="large")
                with c1:
                    comp["comp_id"] = st.text_input(
                        "Component ID (optional)",
                        value=_s(comp.get("comp_id")),
                        placeholder="e.g., Component A",
                        key=_k("comp_id", ci),
                    )
                with c2:
                    comp["title"] = st.text_input(
                        "Component title",
                        value=_s(comp.get("title")),
                        placeholder="e.g., Water Supply System",
                        key=_k("comp_title", ci),
                    )

                st.divider()

                observations: List[Dict[str, Any]] = comp.get("observations") or []
                if not observations:
                    observations = [_ensure_obs_schema({})]

                for oi in range(len(observations)):
                    it = _ensure_obs_schema(observations[oi])

                    raw_title = _obs_title_raw(it)
                    numbered = _numbered_title(SECTION_NO, global_obs_idx, raw_title) if raw_title else ""
                    header_title = numbered if numbered else f"Observation {oi + 1}"

                    with st.container(border=True):
                        hdr = st.columns([3, 1], gap="small")
                        with hdr[0]:
                            st.markdown(f"**{header_title}**")
                            st.caption("A title is required for this observation to appear in the report.")
                        with hdr[1]:
                            st.button(
                                "Remove",
                                use_container_width=True,
                                key=_k("rm_obs", ci, oi),
                                on_click=_remove_observation,
                                args=(ci, oi),
                                disabled=(len(observations) <= 1),
                            )

                        # Audio
                        if audio_url_list:
                            a1, a2 = st.columns([2, 1], gap="small")
                            with a1:
                                cur_audio = _s(it.get("audio_url"))
                                opts = [""] + audio_url_list
                                idx = opts.index(cur_audio) if cur_audio in opts else 0
                                picked_audio = st.selectbox(
                                    "Audio (optional)",
                                    options=opts,
                                    index=idx,
                                    format_func=lambda u: audio_label_by_url.get(u, "None") if u else "None",
                                    key=_k("audio_pick", ci, oi),
                                    help="Select an audio and play it.",
                                )
                                it["audio_url"] = picked_audio
                            with a2:
                                st.caption("Play")
                                if it["audio_url"]:
                                    ab, mime = _fetch_and_cache_audio(it["audio_url"], fetch_audio=fetch_audio)
                                    if ab:
                                        st.audio(ab, format=mime or "audio/aac")
                                    else:
                                        st.audio(it["audio_url"])
                        else:
                            it["audio_url"] = ""

                        st.divider()

                        # Title
                        m1, m2 = st.columns([1, 2], gap="small")
                        with m1:
                            it["title_mode"] = st.radio(
                                "Title type",
                                options=["Select", "Custom"],
                                index=0 if it.get("title_mode") != "Custom" else 1,
                                key=_k("title_mode", ci, oi),
                                horizontal=True,
                            )
                        with m2:
                            if it["title_mode"] == "Select":
                                opts = [""] + DEFAULT_OBSERVATION_TITLES
                                cur = _s(it.get("title_selected"))
                                idx = opts.index(cur) if cur in opts else 0
                                it["title_selected"] = st.selectbox(
                                    "Select title",
                                    options=opts,
                                    index=idx,
                                    key=_k("title_sel", ci, oi),
                                )
                                it["title_custom"] = ""
                            else:
                                it["title_custom"] = st.text_input(
                                    "Custom title",
                                    value=_s(it.get("title_custom")),
                                    placeholder="Write a new title...",
                                    key=_k("title_custom", ci, oi),
                                )
                                it["title_selected"] = ""

                        title_final = _obs_title_raw(it)

                        it["text"] = st.text_area(
                            "Observation text (overall)",
                            value=_s(it.get("text")),
                            height=110,
                            key=_k("obs_text", ci, oi),
                            placeholder="Write the overall observation details here...",
                        )

                        # Photos + notes
                        if not urls:
                            st.info("No photo URLs are available for this record.")
                            it["photos"] = []
                        else:
                            if not title_final:
                                st.warning("Select or enter a title to enable photo selection for this observation.")
                                selected_urls: List[str] = []
                            else:
                                selected_urls = st.multiselect(
                                    "Select photos for this observation",
                                    options=urls,
                                    default=[
                                        p.get("url") for p in (it.get("photos") or [])
                                        if isinstance(p, dict) and p.get("url") in urls
                                    ],
                                    format_func=lambda u: labels.get(u, u),
                                    key=_k("obs_photos", ci, oi),
                                )

                            for u in selected_urls:
                                _fetch_and_cache_image(u, fetch_image=fetch_image)

                            it["photos"] = _normalize_photos(selected_urls, it.get("photos") or [])

                            if it["photos"]:
                                st.markdown("**Photo notes (one note per photo)**")
                                for pj, ph in enumerate(it["photos"]):
                                    u = _s(ph.get("url"))
                                    colA, colB = st.columns([1, 2], gap="small")
                                    with colA:
                                        st.caption(labels.get(u, u))
                                        if show_previews and u:
                                            bts = st.session_state.get(SS_PHOTO_BYTES, {}).get(u)
                                            if bts:
                                                st.image(bts, use_container_width=True)
                                            else:
                                                st.caption("Image bytes are not available.")
                                    with colB:
                                        ph["note"] = st.text_area(
                                            "Note",
                                            value=_s(ph.get("note")),
                                            height=70,
                                            key=_k("photo_note", ci, oi, pj),
                                            placeholder="Write a short note for this photo (will appear in the report).",
                                        )

                        observations[oi] = it

                        # ✅ advance global index only when it has a title
                        if _s(title_final):
                            global_obs_idx += 1

                comp["observations"] = observations
                comps[ci] = comp

        # ---------------------------------------------------------------------
        # ✅ SECOND PASS: build observations_valid with TRUE global numbering
        # ---------------------------------------------------------------------
        photo_bytes_cache: Dict[str, bytes] = st.session_state.get(SS_PHOTO_BYTES, {}) or {}
        global_idx = 1
        for ci in range(len(comps)):
            comp = _ensure_component_schema(comps[ci])
            observations = comp.get("observations") or []
            valid, global_idx = _build_valid_observations_global(
                SECTION_NO,
                observations,
                start_index_1based=global_idx,
                photo_bytes_cache=photo_bytes_cache,
            )
            comp["observations_valid"] = valid
            comps[ci] = comp

        st.session_state[SS_OBS] = comps

        # ✅ this is what the main Tool 6 expects
        st.session_state["tool6_component_observations_final"] = comps

        total_valid = sum(len(c.get("observations_valid") or []) for c in comps if isinstance(c, dict))
        if total_valid > 0:
            status_card("Saved", f"{total_valid} observation(s) will be included in the report.", level="success")
        else:
            status_card("Saved", "No titled observations yet. Add/select a title to include items in the report.", level="warning")

        st.divider()

        # Word-like preview (heavy) — only run if enabled
        if st.session_state.get(_k("show_word_preview")):
            st.markdown("### Preview (Word-like)")

            docx_preview = _get_or_build_obs_preview_docx(ctx, comps, photo_bytes_cache)
            if not docx_preview:
                status_card("Preview failed", "Could not build preview DOCX.", level="error")
            else:
                png = _get_or_build_obs_preview_png(docx_preview)
                if png:
                    st.image(png, use_container_width=False, caption="Observations preview (rendered from DOCX)")
                else:
                    st.info(
                        "Exact Word-like PNG preview needs Windows + MS Word + pywin32 + pymupdf.\n"
                        "You can still download the DOCX preview below."
                    )

                st.download_button(
                    "Download Observations Preview (DOCX)",
                    data=docx_preview,
                    file_name=f"Tool6_ObservationsPreview_{ctx.tpm_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        card_close()

    return total_valid > 0
