# src/report_sections/conclusion.py
from __future__ import annotations

import re
from typing import Any, Dict, List, Optional, Union

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph


# ============================================================
# Inline helpers (replacing src.report_sections._word_common)
# ============================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def set_run(
    run,
    font: str,
    size: Union[int, float],
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    """Safe run styling with eastAsia font set."""
    try:
        run.font.name = font
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass
    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass
    run.bold = bool(bold)
    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: Union[int, float] = 0,
    after_pt: Union[int, float] = 0,
    line_spacing: float = 1.0,
) -> None:
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def body(
    doc: Document,
    text: str,
    *,
    font: str = "Times New Roman",
    size: int = 11,
    bold: bool = False,
    color: Optional[RGBColor] = None,
    line_spacing: float = 1.15,
    after_pt: float = 0,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> Paragraph:
    """
    Add a standard body paragraph with consistent formatting.
    """
    p = doc.add_paragraph()
    tight_paragraph(p, align=align, before_pt=0, after_pt=float(after_pt), line_spacing=line_spacing)
    r = p.add_run(s(text))
    set_run(r, font, size, bold=bold, color=color)
    return p


def bullets_from_text(text: str) -> List[str]:
    """
    Split free text into bullet items.
    Handles:
      - newline-separated
      - lines starting with -, *, •
      - semicolon-separated fallback
    """
    t = s(text)
    if not t:
        return []

    # Normalize newlines
    lines = [ln.strip() for ln in t.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    lines = [ln for ln in lines if ln]

    items: List[str] = []

    if len(lines) > 1:
        for ln in lines:
            ln = re.sub(r"^\s*[-•*]\s*", "", ln).strip()
            if ln:
                items.append(ln)
        return items

    # Single-line fallback: split by ; or •
    one = lines[0] if lines else t
    # If it already contains bullet markers, split on them
    if "•" in one:
        parts = [p.strip() for p in one.split("•") if p.strip()]
        return parts

    # Semicolon split fallback
    if ";" in one:
        parts = [p.strip() for p in one.split(";") if p.strip()]
        return parts

    return []


def add_heading(
    doc: Document,
    text: str,
    *,
    level: int = 1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font: str = "Cambria",
    size: int = 16,
    bold: bool = True,
    color: Optional[RGBColor] = None,
) -> Paragraph:
    """
    Add a REAL Word heading paragraph (TOC-safe).
    """
    lvl = max(1, min(int(level), 9))
    p = doc.add_paragraph()
    try:
        p.style = f"Heading {lvl}"
    except Exception:
        pass

    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1.0)
    r = p.add_run(s(text))
    set_run(r, font, size, bold=bold, color=color)
    return p


def set_paragraph_bottom_border(
    paragraph: Paragraph,
    *,
    color_hex: str,
    size_eighths: int = 12,
    space: int = 2,
) -> None:
    """
    Add/Update an orange bottom border under a paragraph (idempotent).
    size_eighths: Word uses 1/8 pt units. 12 => 1.5pt.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), s(color_hex).replace("#", ""))


# -------------------------
# Section style (match your new rule)
# -------------------------
TITLE_FONT = "Cambria"
TITLE_SIZE = 16  # ✅ Heading 1 titles = 16
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"


def _heading_with_orange_line(
    doc: Document,
    text: str,
    *,
    level: int,
    font: str = TITLE_FONT,
    size: int = TITLE_SIZE,
    color: Optional[RGBColor] = TITLE_BLUE,
    orange_hex: str = ORANGE_HEX,
    after_pt: float = 6,
) -> None:
    """
    Create a REAL Word Heading (TOC-safe) and add an orange bottom border under it.
    """
    p = add_heading(
        doc,
        s(text),
        level=int(level),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font=font,
        size=int(size),
        bold=True,
        color=color,
    )

    # consistent spacing
    tight_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before_pt=0,
        after_pt=float(after_pt),
        line_spacing=1,
    )

    # orange underline (bottom border)
    set_paragraph_bottom_border(
        p,
        color_hex=orange_hex,
        size_eighths=12,  # 12 => 1.5pt
        space=2,
    )


def add_conclusion_section(
    doc: Document,
    *,
    conclusion_text: Optional[str] = None,
    key_points: Optional[List[str]] = None,
    recommendations_summary: Optional[str] = None,
    section_no: str = "7",
) -> None:
    """
    Final section: Conclusion (TOC-safe)
    - Heading 1 for main title (shows in TOC)
    - Heading 2 for subsections (shows in TOC when TOC levels include 1-3)
    """
    doc.add_page_break()

    # ✅ Heading 1 + size 16 + orange underline
    _heading_with_orange_line(
        doc,
        f"{s(section_no)}.        Conclusion:",
        level=1,
        after_pt=6,
    )

    # 1) Main conclusion paragraph (always print something)
    main_text = s(conclusion_text)
    if main_text:
        body(doc, main_text)
    else:
        body(
            doc,
            "Overall, the monitoring confirmed that the assessed WASH intervention is functional and "
            "providing services to the beneficiary community. Addressing the observed technical and "
            "operational gaps through timely corrective actions and strengthened O&M capacity will "
            "improve system reliability and long-term sustainability."
        )

    # small spacing after main paragraph (controlled)
    doc.add_paragraph("")

    # 2) Key points (optional)
    kp = [s(x) for x in (key_points or []) if s(x)]
    if kp:
        add_heading(
            doc,
            "Key Points",
            level=2,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font="Cambria",
            size=12,
            bold=True,
            color=None,  # black
        )
        tight_paragraph(doc.paragraphs[-1], before_pt=0, after_pt=4, line_spacing=1)

        for it in kp:
            body(doc, f"• {it}", align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)

        doc.add_paragraph("")

    # 3) Recommendations summary (optional)
    rec_text = s(recommendations_summary)
    if rec_text:
        add_heading(
            doc,
            "Recommendations Summary",
            level=2,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font="Cambria",
            size=12,
            bold=True,
            color=None,  # black
        )
        tight_paragraph(doc.paragraphs[-1], before_pt=0, after_pt=4, line_spacing=1)

        items = bullets_from_text(rec_text)
        if items:
            for it in items:
                body(doc, f"• {it}", align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
        else:
            body(doc, rec_text)

        doc.add_paragraph("")


# ✅ Compatibility wrapper (some builders may call add_conclusion)
def add_conclusion(
    doc: Document,
    *,
    row: Optional[Dict[str, Any]] = None,
    conclusion_text: Optional[str] = None,
    key_points: Optional[List[str]] = None,
    recommendations_summary: Optional[str] = None,
    component_observations: Optional[List[Dict[str, Any]]] = None,
    severity_by_no: Any = None,
    section_no: str = "7",
    **kwargs,
) -> None:
    """
    Compatibility wrapper so existing pipelines keep working.
    Extra params are accepted but not required here.
    """
    _ = row, component_observations, severity_by_no, kwargs

    add_conclusion_section(
        doc,
        conclusion_text=conclusion_text,
        key_points=key_points,
        recommendations_summary=recommendations_summary,
        section_no=section_no,
    )
