# src/report_sections/data_collection_methods.py
from __future__ import annotations

from typing import Any, Dict, Optional, List, Union

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph


# ============================================================
# Inline helpers (replacing src.report_sections._word_common)
# ============================================================
def s(v: Any) -> str:
    """Safe stringify + trim."""
    return "" if v is None else str(v).strip()


def set_run(
    run,
    font: str,
    size: Union[int, float],
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    """
    Safe run styling with eastAsia font set
    (prevents Arabic/Farsi/Unicode fallback rendering issues).
    """
    try:
        run.font.name = font
    except Exception:
        pass

    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass

    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass

    run.bold = bool(bold)

    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: Union[int, float] = 0,
    after_pt: Union[int, float] = 0,
    line_spacing: float = 1.0,
) -> None:
    """Standard paragraph formatting (tight, predictable)."""
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def parse_bool_like(v: Any) -> Optional[bool]:
    """
    Returns True/False/None.
    Accepts: bool, 0/1, yes/no, checked/unchecked, ✅/❌, etc.
    """
    if v is None:
        return None
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        iv = int(v)
        if iv == 1:
            return True
        if iv == 0:
            return False

    sv = s(v).strip().lower()
    if sv in {"yes", "y", "true", "1", "checked", "✔", "✅"}:
        return True
    if sv in {"no", "n", "false", "0", "unchecked", "✘", "❌"}:
        return False
    return None


def add_heading(
    doc: Document,
    text: str,
    *,
    level: int = 1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font: str = "Cambria",
    size: int = 14,
    bold: bool = True,
    color: Optional[RGBColor] = None,
) -> Paragraph:
    """
    Add a REAL Word Heading (Heading 1/2/3/...)
    REQUIRED for Table of Contents to work correctly.
    """
    lvl = max(1, min(int(level), 9))

    p = doc.add_paragraph(s(text))
    try:
        p.style = f"Heading {lvl}"
    except Exception:
        pass

    p.alignment = align

    # style runs
    for r in p.runs:
        r.bold = bool(bold)
        try:
            r.font.name = font
        except Exception:
            pass
        try:
            r._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
        except Exception:
            pass
        try:
            r.font.size = Pt(int(size))
        except Exception:
            pass
        if color is not None:
            try:
                r.font.color.rgb = color
            except Exception:
                pass

    return p


def set_paragraph_bottom_border(
    paragraph: Paragraph,
    *,
    color_hex: str,
    size_eighths: int = 12,
    space: int = 2,
) -> None:
    """Adds/updates a bottom border to a paragraph (idempotent)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), s(color_hex).replace("#", ""))


def add_section_title_h1(
    doc: Document,
    text: str,
    *,
    font: str = "Cambria",
    size: Union[int, float] = 16,
    color: RGBColor = RGBColor(0, 112, 192),
    orange_hex: str = "ED7D31",
    after_pt: Union[int, float] = 6,
    border_size_eighths: int = 12,
    border_space: int = 2,
) -> Paragraph:
    """
    Standard section title:
      - Real Heading 1 (TOC-friendly)
      - Font size 16
      - Blue title + orange underline (on SAME paragraph)
    """
    p = add_heading(
        doc,
        s(text),
        level=1,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font=font,
        size=int(size),
        bold=True,
        color=color,
    )
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1)
    set_paragraph_bottom_border(
        p,
        color_hex=orange_hex,
        size_eighths=int(border_size_eighths),
        space=int(border_space),
    )
    return p


# -------------------------
# Section style
# -------------------------
TITLE_TEXT = "3.        Data Collection Methods:"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"

BODY_FONT = "Times New Roman"
BODY_SIZE = 11

AFTER_SECTION_GAP_PT = 24


# -------------------------
# Internal helpers
# -------------------------
def _yes(v: Any) -> bool:
    return parse_bool_like(v) is True


def _pick(row: Dict[str, Any], overrides: Dict[str, Any], *keys: str) -> Any:
    for k in keys:
        if k in overrides and overrides.get(k) not in (None, "", " "):
            return overrides.get(k)
    for k in keys:
        if row.get(k) not in (None, "", " "):
            return row.get(k)
    return None


def _add_numbered_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1.0)
    set_run(p.add_run(s(text)), BODY_FONT, BODY_SIZE, bold=False)


def _add_two_line_gap(doc: Document) -> None:
    p = doc.add_paragraph("")
    tight_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before_pt=0,
        after_pt=float(AFTER_SECTION_GAP_PT),
        line_spacing=1.0,
    )


def _split_lines_to_items(text: str) -> List[str]:
    lines = [ln.strip() for ln in s(text).splitlines()]
    return [ln for ln in lines if ln]


# -------------------------
# Main section
# -------------------------
def add_data_collection_methods(
    doc: Document,
    row: Dict[str, Any],
    overrides: Optional[Dict[str, Any]] = None,
) -> None:
    row = row or {}
    overrides = overrides or {}

    doc.add_page_break()

    add_section_title_h1(
        doc,
        TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=TITLE_BLUE,
        orange_hex=ORANGE_HEX,
        after_pt=6,
    )

    # =========================================================
    # ✅ NEW: If user provided finalized text from Step 7, use it
    # =========================================================
    manual_list_text = s(overrides.get("D_methods_list_text"))
    manual_narrative = s(overrides.get("D_methods_narrative_text"))

    if manual_list_text or manual_narrative:
        items = _split_lines_to_items(manual_list_text) if manual_list_text else []
        if not items:
            items = [
                "The monitoring visit applied standard Third-Party Monitoring (TPM) data collection techniques in line with UNICEF WASH guidelines."
            ]

        for it in items:
            _add_numbered_item(doc, it)

        p_gap = doc.add_paragraph("")
        tight_paragraph(p_gap, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=6, line_spacing=1)

        narrative = manual_narrative.strip() if manual_narrative else (
            "The Third-Party Monitoring (TPM) assessment was conducted using a structured mixed-methods approach..."
        )
        p = doc.add_paragraph()
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before_pt=0, after_pt=0, line_spacing=1.15)
        set_run(p.add_run(narrative), BODY_FONT, BODY_SIZE, bold=False)

        _add_two_line_gap(doc)
        return

    # =========================================================
    # 1) Detect available evidence (prefer overrides, fallback row)
    # =========================================================
    has_contract = _yes(_pick(row, overrides, "D1_contract_available"))
    has_journal = _yes(_pick(row, overrides, "D1_journal_available"))
    has_boq = _yes(_pick(row, overrides, "D2_boq_available"))
    has_drawings = _yes(_pick(row, overrides, "D2_drawings_available"))
    has_geo_tests = _yes(_pick(row, overrides, "D3_geophysical_tests_available"))
    has_wq_tests = _yes(_pick(row, overrides, "D4_water_quality_tests_available"))
    has_pump_tests = _yes(_pick(row, overrides, "D4_pump_test_results_available"))

    has_observation = _yes(_pick(row, overrides, "D0_direct_observation"))
    has_interview = _yes(_pick(row, overrides, "D0_key_informant_interview"))
    has_photos = _yes(_pick(row, overrides, "D0_photos_taken"))
    has_gps = _yes(_pick(row, overrides, "D0_gps_points_recorded"))

    # =========================================================
    # 2) Build DOCUMENT REVIEW text dynamically
    # =========================================================
    reviewed_docs: List[str] = []

    if has_boq:
        reviewed_docs.append("Bill of Quantities (BOQ)")
    if has_drawings:
        reviewed_docs.append("approved technical drawings")
    if has_contract:
        reviewed_docs.append("contract documents")
    if has_journal:
        reviewed_docs.append("site journal and progress records")
    if has_geo_tests:
        reviewed_docs.append("geophysical and hydrological test reports")
    if has_wq_tests:
        reviewed_docs.append("water quality test results")
    if has_pump_tests:
        reviewed_docs.append("pump test results")

    doc_review_phrase = ""
    if reviewed_docs:
        doc_review_phrase = "Review of project documentation, including " + ", ".join(reviewed_docs) + "."

    # =========================================================
    # 3) Build numbered METHODS list (automatic)
    # =========================================================
    methods: List[str] = []

    if has_observation:
        methods.append("Direct technical observation of work progress and construction quality on-site.")

    if doc_review_phrase:
        methods.append(doc_review_phrase)

    if has_interview:
        methods.append(
            "Semi-structured interviews with technical staff of the contracted company, implementing partner personnel, "
            "and Community Development Council (CDC) members."
        )

    if has_photos:
        methods.append(
            "Collection and review of geo-referenced photographic evidence to verify physical progress and workmanship."
        )

    if has_gps:
        methods.append(
            "Verification of GPS coordinates and location data to confirm site positioning and component alignment."
        )

    if not methods:
        methods.append(
            "The monitoring visit applied standard Third-Party Monitoring (TPM) data collection techniques in line with UNICEF WASH guidelines."
        )

    for m in methods:
        _add_numbered_item(doc, m)

    p_gap = doc.add_paragraph("")
    tight_paragraph(p_gap, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=6, line_spacing=1)

    narrative = (
        "The Third-Party Monitoring (TPM) assessment was conducted using a structured mixed-methods approach, combining "
        "direct on-site technical observation, systematic review of available project documentation, and qualitative "
        "engagement with relevant stakeholders. The monitoring focused on verifying construction quality, system "
        "functionality, and compliance with approved designs and contractual requirements, while identifying technical "
        "and operational risks that may affect performance and sustainability. Physical and documentary evidence was "
        "assessed across all applicable project components, and findings were analyzed, categorized by severity, and "
        "linked to practical corrective actions in accordance with UNICEF WASH standards and third-party monitoring protocols."
    )

    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before_pt=0, after_pt=0, line_spacing=1.15)
    set_run(p.add_run(narrative), BODY_FONT, BODY_SIZE, bold=False)

    _add_two_line_gap(doc)
