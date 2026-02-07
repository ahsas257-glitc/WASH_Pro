# src/report_sections/executive_summary.py
from __future__ import annotations

from datetime import date, datetime
from typing import Any, Dict, Optional, List, Iterable, Union

from docx.document import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# Inline helpers (replacing src.report_sections._word_common)
# ============================================================
def s(v: Any) -> str:
    """Safe stringify + trim."""
    return "" if v is None else str(v).strip()


def _is_nonempty(v: Any) -> bool:
    sv = s(v)
    return sv not in ("", " ")


def pick_first_nonempty(
    row: Dict[str, Any],
    overrides: Dict[str, Any],
    keys: Union[Iterable[str], str],
) -> Any:
    """
    Return first non-empty value from overrides then row, scanning provided keys.
    keys can be a list/tuple of key names, or a single string key.
    """
    if isinstance(keys, str):
        keys_list = [keys]
    else:
        keys_list = list(keys)

    for k in keys_list:
        if k in overrides and _is_nonempty(overrides.get(k)):
            return overrides.get(k)
    for k in keys_list:
        if _is_nonempty(row.get(k)):
            return row.get(k)
    return None


def parse_bool_like(v: Any) -> Optional[bool]:
    """Returns True/False/None for common bool-like inputs."""
    if v is None:
        return None
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        iv = int(v)
        if iv == 1:
            return True
        if iv == 0:
            return False

    sv = s(v).lower()
    if sv in {"yes", "y", "true", "1", "checked", "✔", "✅"}:
        return True
    if sv in {"no", "n", "false", "0", "unchecked", "✘", "❌"}:
        return False
    return None


def as_yes(v: Any) -> bool:
    return parse_bool_like(v) is True


def as_no(v: Any) -> bool:
    return parse_bool_like(v) is False


def norm_phrase(v: Any) -> str:
    """
    Normalize free-text phrases:
    - stringify
    - trim
    - collapse internal whitespace
    - remove trailing punctuation-only tails
    """
    sv = s(v)
    if not sv:
        return ""
    sv = " ".join(sv.split())
    sv = sv.rstrip(" .;:,")
    return sv


def date_only_isoish(v: Any) -> str:
    """
    Convert common date/datetime/ISO-ish strings to YYYY-MM-DD.
    If parsing fails, returns best-effort date-like token.
    """
    if v is None:
        return ""

    if isinstance(v, datetime):
        return v.date().isoformat()
    if isinstance(v, date):
        return v.isoformat()

    sv = s(v)
    if not sv:
        return ""

    # Clean ISO-ish: remove milliseconds and Z
    sv_clean = sv.strip().replace("T", " ").replace("Z", "")
    if "." in sv_clean:
        sv_clean = sv_clean.split(".")[0].strip()

    # Try parse common formats
    for fmt in (
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
        "%d/%m/%Y",
        "%m/%d/%Y",
    ):
        try:
            dt = datetime.strptime(sv_clean, fmt)
            return dt.date().isoformat()
        except Exception:
            pass

    # Best effort: take first token if it looks like a date
    first = sv_clean.split(" ")[0].strip()
    return first


def build_location_phrase(village: str, district: str, province: str) -> str:
    """
    Build a clean location phrase like: "Village, District, Province".
    Returns "" if nothing is available.
    """
    parts = [p for p in [s(village), s(district), s(province)] if p]
    return ", ".join(parts)


def add_body_paragraph(
    doc: Document,
    text: str,
    *,
    font: str = "Times New Roman",
    size: int = 11,
    line_spacing: float = 1.0,
    before_pt: float = 0,
    after_pt: float = 0,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    """
    Add a body paragraph with stable spacing and font.
    """
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)

    run = p.add_run(s(text))
    try:
        run.font.name = font
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass
    try:
        run.font.size = Pt(int(size))
    except Exception:
        pass


# -------------------------
# Section-specific style only
# -------------------------
TITLE_TEXT = "2.    Executive Summary:"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16  # ✅ requested

BODY_FONT = "Times New Roman"
BODY_SIZE = 11


def _add_heading1_with_orange_line(
    doc: Document,
    *,
    text: str,
    font: str = "Cambria",
    size: int = 16,
    orange_hex: str = "F28C28",
    line_sz_eights: str = "12",  # Word border size in 1/8 pt (12 => 1.5pt)
) -> None:
    """
    Adds a Heading 1 title with an orange bottom border line.
    - Title paragraph: style = 'Heading 1'
    - Font size: configurable (default 16)
    - Orange line: paragraph bottom border
    """
    p = doc.add_paragraph()
    try:
        p.style = "Heading 1"
    except Exception:
        # If style not found, still proceed with direct formatting
        pass

    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)

    # Ensure paragraph formatting is tight/clean (optional but safe)
    try:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
    except Exception:
        pass

    # Add orange bottom border
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(line_sz_eights))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), orange_hex)


def add_executive_summary(
    doc: Document,
    row: Dict[str, Any],
    overrides: Optional[Dict[str, Any]] = None,
) -> None:
    """
    Section 2: Executive Summary (Streamlit-friendly)
    - Only section logic here
    - ✅ No dependency on _word_common
    """
    overrides = overrides or {}
    row = row or {}

    doc.add_page_break()

    # ✅ Title in Heading 1 + ✅ font size 16 + ✅ orange line under
    _add_heading1_with_orange_line(
        doc,
        text=TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        orange_hex="F28C28",
        line_sz_eights="12",
    )

    # -------------------------
    # Robust extraction (single-pass)
    # -------------------------
    province = s(pick_first_nonempty(row, overrides, ["A01_Province", "province", "Province"]))
    district = s(pick_first_nonempty(row, overrides, ["A02_District", "district", "District"]))
    village = s(pick_first_nonempty(row, overrides, ["Village", "village", "Community"]))
    project_name = s(pick_first_nonempty(row, overrides, ["Activity_Name", "project", "Project_Name"]))

    visit_date = date_only_isoish(pick_first_nonempty(row, overrides, ["starttime", "visit_date", "Date_of_Visit"]))

    status_raw = norm_phrase(pick_first_nonempty(row, overrides, ["Project_Status", "project_status", "status"]))
    progress_raw = norm_phrase(pick_first_nonempty(row, overrides, ["Project_progress", "project_progress", "progress"]))

    pipeline_issue = pick_first_nonempty(row, overrides, ["pipeline_installation_issue", "pipeline_issue"])
    leakage = pick_first_nonempty(row, overrides, ["leakage_observed", "leakage"])
    dust_panels = pick_first_nonempty(row, overrides, ["solar_panel_dust", "dust_panels"])
    training = pick_first_nonempty(row, overrides, ["community_training_conducted", "training_conducted"])

    location = build_location_phrase(village, district, province) or "the monitored location"

    # -------------------------
    # Paragraph 1 – Purpose
    # -------------------------
    proj_phrase = (
        f"the Solar Water Supply project with household connections ({project_name})"
        if project_name
        else "the Solar Water Supply project with household connections"
    )
    date_phrase = f" on {visit_date}" if visit_date else ""

    add_body_paragraph(
        doc,
        "This Third-Party Monitoring (TPM) field visit was conducted to assess the technical "
        f"implementation, functionality, and compliance of {proj_phrase} in {location}{date_phrase}. "
        "The visit focused on verifying system operational status, adherence to approved designs "
        "and Bill of Quantities (BoQ), and identifying any technical or operational risks that may "
        "affect long-term system performance.",
        font=BODY_FONT,
        size=BODY_SIZE,
        line_spacing=1.0,
        after_pt=0,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    # -------------------------
    # Paragraph 1.1 – Status/Progress (optional)
    # -------------------------
    sp_bits: List[str] = []
    if status_raw:
        sp_bits.append(f"Project status was reported as {status_raw}.")
    if progress_raw:
        sp_bits.append(f"Overall progress was reported as {progress_raw}.")
    if sp_bits:
        add_body_paragraph(
            doc,
            " ".join(sp_bits),
            font=BODY_FONT,
            size=BODY_SIZE,
            line_spacing=1.0,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    # -------------------------
    # Paragraph 2 – System status (standard)
    # -------------------------
    add_body_paragraph(
        doc,
        "The assessment confirmed that the water supply system infrastructure—including bore wells, "
        "solar-powered pumping system, reservoirs, boundary wall, guard room, latrine, and stand taps—"
        "has been constructed and is currently operational. The system is supplying water to the "
        "targeted community, and the majority of stand taps were observed to be functional at the "
        "time of the visit.",
        font=BODY_FONT,
        size=BODY_SIZE,
        line_spacing=1.0,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    # -------------------------
    # Paragraph 3 – Issues (dynamic)
    # -------------------------
    issues: List[str] = []
    if as_yes(pipeline_issue):
        issues.append("pipeline installation and protection deficiencies")
    if as_yes(leakage):
        issues.append("localized leakages in the distribution network")
    if as_yes(dust_panels):
        issues.append("reduced solar panel efficiency due to dust accumulation")
    if as_no(training):
        issues.append("lack of formal community training on system operation and maintenance")

    if issues:
        add_body_paragraph(
            doc,
            "However, several technical and operational gaps were identified during the monitoring. "
            "These include " + ", ".join(issues) +
            ". While minor construction defects were observed in selected concrete works, no critical "
            "structural failures were noted during the visit.",
            font=BODY_FONT,
            size=BODY_SIZE,
            line_spacing=1.0,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
    else:
        add_body_paragraph(
            doc,
            "No major technical or operational deficiencies were identified during the monitoring, "
            "and the system generally complies with the approved technical specifications.",
            font=BODY_FONT,
            size=BODY_SIZE,
            line_spacing=1.0,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )

    # -------------------------
    # Paragraph 4 – Conclusion
    # -------------------------
    add_body_paragraph(
        doc,
        "Overall, the project is functional and delivering water services to the beneficiary community. "
        "Addressing the identified gaps through timely corrective actions and strengthening community "
        "capacity will further enhance system reliability, operational safety, and long-term "
        "sustainability of the water supply service.",
        font=BODY_FONT,
        size=BODY_SIZE,
        line_spacing=1.0,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
