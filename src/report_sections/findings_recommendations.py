from __future__ import annotations

import io
from typing import Any, Dict, List, Optional, Tuple

from docx.document import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


# ============================================================
# Major Findings table — exact widths (inches)
# ============================================================
COL_NO = 0.38
COL_FIND = 2.88
COL_COMP = 0.91
COL_PHOTO = 3.40

PHOTO_STACK_WIDTH_IN = 3.10  # fits inside COL_PHOTO

# Border thickness: 1/4pt => 0.25pt * 8 = 2 (Word uses 1/8 pt units)
BORDER_SZ_EIGHTHS = "2"


# ============================================================
# Utils
# ============================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def _compact(p) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _one_line_gap(doc: Document) -> None:
    p = doc.add_paragraph("")
    _compact(p)


def _clear_cell(cell) -> None:
    cell.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = ""
        _compact(cell.paragraphs[0])


def _set_table_fixed_layout(tbl) -> None:
    """
    Prevent Word auto-resize.
    Uses tblLayout=fixed and disables autofit.
    """
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _shade_cell(cell, fill_hex: str) -> None:
    fill_hex = s(fill_hex).replace("#", "") or "FFFFFF"
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_table_borders(tbl, *, color_hex: str = "000000", size_eighths: str = BORDER_SZ_EIGHTHS) -> None:
    """
    Apply borders to whole table (including inside borders).
    size_eighths: Word uses 1/8pt units. For 1/4pt => "2".
    """
    color_hex = s(color_hex).replace("#", "") or "000000"

    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size_eighths))   # ✅ 1/4pt => 2
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def _apply_table_col_widths_exact(tbl, widths_in: List[float]) -> None:
    """
    Force exact column widths (docx + tcW dxa) to stop Word reflow.
    """
    _set_table_fixed_layout(tbl)

    for i, w in enumerate(widths_in):
        tbl.columns[i].width = Inches(float(w))
        for cell in tbl.columns[i].cells:
            cell.width = Inches(float(w))
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(int(float(w) * 1440)))  # inches->twips


def _align_cell_center(cell) -> None:
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        pass
    if cell.paragraphs:
        p = cell.paragraphs[0]
        _compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _align_cell_left_top(cell) -> None:
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    except Exception:
        pass
    if cell.paragraphs:
        p = cell.paragraphs[0]
        _compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _clean_png(b: bytes) -> Optional[bytes]:
    """
    Normalize image bytes -> PNG (RGB) to avoid docx rendering issues.
    """
    if not b:
        return None
    try:
        img = Image.open(io.BytesIO(b))
        img = img.convert("RGB")
        out = io.BytesIO()
        img.save(out, format="PNG", optimize=True)
        return out.getvalue()
    except Exception:
        return None


def _add_picture_stack_in_cell(cell, images: List[bytes], *, width_in: float = PHOTO_STACK_WIDTH_IN) -> None:
    """
    Stack multiple images (each in its own right-aligned paragraph).
    """
    _clear_cell(cell)
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    except Exception:
        pass

    if not images:
        return

    for img_bytes in images:
        clean = _clean_png(img_bytes)
        if not clean:
            continue
        p = cell.add_paragraph("")
        _compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(io.BytesIO(clean), width=Inches(float(width_in)))


def _pick_compliance(rdata: Dict[str, Any]) -> str:
    """
    Ensure compliance prints EXACTLY what user selected.
    Accept both keys to be robust:
      - "Compliance" (your Step4 UI)
      - "compliance" (some older code)
    """
    v = rdata.get("Compliance")
    if s(v):
        return s(v)
    return s(rdata.get("compliance"))


def _extract_images_for_row(
    rdata: Dict[str, Any],
    photo_bytes: Dict[str, bytes],
) -> List[bytes]:
    """
    STRICT priority to guarantee annotated marks show in the Word report.

    Supports BOTH schemas:
      A) New Step4 (recommended):
         - photo_bytes (single bytes; annotated already preferred in merge)
         - photo (url)
      B) Older schemas:
         - annotated_photo_bytes_list / photo_bytes_list / photos[list] / photo

    Priority:
      1) rdata["photo_bytes"] (single bytes)  ✅ best, simplest
      2) annotated_photo_bytes_list           ✅ explicit annotated list
      3) photo_bytes_list                     ✅ explicit original list
      4) photos urls -> photo_bytes dict
      5) photo url -> photo_bytes dict
    """
    out: List[bytes] = []

    # 1) single bytes
    b0 = rdata.get("photo_bytes")
    if isinstance(b0, (bytes, bytearray)) and b0:
        out.append(bytes(b0))

    # 2) annotated list
    ann_list = rdata.get("annotated_photo_bytes_list")
    if isinstance(ann_list, list):
        for b in ann_list:
            if isinstance(b, (bytes, bytearray)) and b:
                out.append(bytes(b))

    # 3) original list
    orig_list = rdata.get("photo_bytes_list")
    if isinstance(orig_list, list):
        for b in orig_list:
            if isinstance(b, (bytes, bytearray)) and b:
                out.append(bytes(b))

    # 4) urls list
    photos = rdata.get("photos")
    if isinstance(photos, list):
        for u in photos:
            u = s(u)
            if not u:
                continue
            b = photo_bytes.get(u)
            if isinstance(b, (bytes, bytearray)) and b:
                out.append(bytes(b))

    # 5) single url
    u1 = s(rdata.get("photo"))
    if u1:
        b = photo_bytes.get(u1)
        if isinstance(b, (bytes, bytearray)) and b:
            out.append(bytes(b))

    # soft dedup
    dedup: List[bytes] = []
    seen = set()
    for b in out:
        sig = (len(b), b[:32])
        if sig in seen:
            continue
        seen.add(sig)
        dedup.append(b)

    return dedup


def _add_major_findings_table(
    doc: Document,
    *,
    rows: List[Dict[str, Any]],
    photo_bytes: Dict[str, bytes],
) -> None:
    """
    Build one Major Findings table with:
      - exact widths
      - black borders 1/4 pt
      - header blue
      - compliance EXACT from user selection
      - photos from annotated bytes first
    """
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    _set_table_fixed_layout(tbl)
    _apply_table_col_widths_exact(tbl, [COL_NO, COL_FIND, COL_COMP, COL_PHOTO])
    _set_table_borders(tbl, color_hex="000000", size_eighths=BORDER_SZ_EIGHTHS)

    # Header
    header_fill_blue = "2F5597"
    hdr_cells = tbl.rows[0].cells
    hdr_texts = ["NO", "Findings", "Compliance", "Photos"]

    for c, txt in zip(hdr_cells, hdr_texts):
        _clear_cell(c)
        _shade_cell(c, header_fill_blue)

        p = c.paragraphs[0]
        _compact(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            pass

        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for idx, rdata in enumerate(rows or [], start=1):
        if not isinstance(rdata, dict):
            continue

        row_cells = tbl.add_row().cells

        # --- NO
        row_cells[0].text = str(idx)
        _align_cell_center(row_cells[0])

        # --- Findings
        row_cells[1].text = s(rdata.get("finding"))
        _align_cell_left_top(row_cells[1])

        # --- Compliance (exact)
        row_cells[2].text = _pick_compliance(rdata)
        _align_cell_center(row_cells[2])

        # --- Photos (annotated-first)
        images = _extract_images_for_row(rdata, photo_bytes)
        _add_picture_stack_in_cell(row_cells[3], images, width_in=PHOTO_STACK_WIDTH_IN)


# =================================================
# PAGE — FINDINGS & RECOMMENDATIONS (INSIDE OBSERVATIONS)
# =================================================
def add_findings_recommendations_page(
    doc: Document,
    *,
    component_observations: List[Dict[str, Any]],
    photo_bytes: Dict[str, bytes],
) -> None:
    """
    Guarantees:
      - user-entered texts are printed (finding + recommendations)
      - compliance prints EXACTLY the selected value
      - annotated image prints EXACTLY (photo_bytes first)
      - borders are exactly 1/4pt (w:sz=2)

    Expected (recommended) from Step4 merge:
      major_table row = {
        "finding": "...",
        "Compliance": "Yes/No/N/A/...",
        "photo": "url",
        "photo_bytes": bytes  # annotated preferred, else original
      }

    Backward compatibility: also supports older lists.
    """
    if not component_observations:
        return

    photo_bytes = photo_bytes or {}

    obs_global_idx = 0  # 5.1, 5.2, ... across components

    for comp in component_observations:
        if not isinstance(comp, dict):
            continue

        for obs in (comp.get("observations_valid") or []):
            if not isinstance(obs, dict):
                continue

            obs_global_idx += 1
            obs_no = f"5.{obs_global_idx}"

            # ----------------------------
            # 5.x.1 Major findings
            # ----------------------------
            mf_title = doc.add_paragraph(f"{obs_no}.1 Major findings:")
            mf_title.style = "Heading 3"
            _compact(mf_title)
            # keep Heading 3 “natural” color; just ensure black if style is weird
            if mf_title.runs:
                mf_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            _one_line_gap(doc)

            rows = obs.get("major_table") or []
            if isinstance(rows, list) and rows:
                _add_major_findings_table(doc, rows=rows, photo_bytes=photo_bytes)
            else:
                p = doc.add_paragraph("No major findings were provided.")
                _compact(p)

            # ----------------------------
            # 5.x.2 Recommendations
            # ----------------------------
            recs = obs.get("recommendations") or []
            if isinstance(recs, list) and any(s(x) for x in recs):
                _one_line_gap(doc)

                rt = doc.add_paragraph(f"{obs_no}.2 Recommendations:")
                rt.style = "Heading 3"
                _compact(rt)
                if rt.runs:
                    rt.runs[0].font.color.rgb = RGBColor(0, 0, 0)

                _one_line_gap(doc)

                for txt in recs:
                    t = s(txt)
                    if not t:
                        continue
                    rp = doc.add_paragraph(t, style="List Bullet")
                    _compact(rp)

            _one_line_gap(doc)
