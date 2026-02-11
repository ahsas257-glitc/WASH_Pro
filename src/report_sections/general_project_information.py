# src/report_sections/general_project_information.py
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional, Tuple, Union

from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# ============================================================
# Small core helpers (string + formatting)
# ============================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def na(v: Any) -> str:
    sv = s(v)
    return sv if sv else "N/A"


def _is_nonempty(v: Any) -> bool:
    return bool(s(v))


# ============================================================
# Bool-like parsing and checkbox rendering
# ============================================================
def parse_bool_like(v: Any) -> Optional[bool]:
    if v is None:
        return None
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        iv = int(v)
        if iv == 1:
            return True
        if iv == 0:
            return False

    sv = s(v).lower()
    if sv in {"yes", "y", "true", "1", "checked", "✔", "✅"}:
        return True
    if sv in {"no", "n", "false", "0", "unchecked", "✘", "❌"}:
        return False
    return None


def _truthy_doc(v: Any) -> bool:
    b = parse_bool_like(v)
    if b is True:
        return True
    if b is False:
        return False
    return _is_nonempty(v)


def _checkbox(checked: bool) -> str:
    return "☒" if checked else "☐"


def three_option_checkbox_line(value: Any, opt1: str, opt2: str, opt3: str) -> str:
    sv = s(value).lower()

    def match(opt: str) -> bool:
        o = opt.lower().strip()
        return sv == o or (sv and o in sv)

    return f"{_checkbox(match(opt1))} {opt1}   {_checkbox(match(opt2))} {opt2}   {_checkbox(match(opt3))} {opt3}"


# ============================================================
# Word layout helpers
# ============================================================
def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: float = 0,
    after_pt: float = 0,
    line_spacing: float = 1.0,
) -> None:
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def set_run(run, font: str, size: Union[int, float], bold: bool = False, color: Optional[RGBColor] = None) -> None:
    try:
        run.font.name = font
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass
    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass
    run.bold = bool(bold)
    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


def _set_paragraph_bottom_border(paragraph: Paragraph, *, color_hex: str, size_eighths: int = 12, space: int = 2) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), s(color_hex).replace("#", ""))


def add_section_title_h1(
    doc: Document,
    text: str,
    *,
    font: str = "Cambria",
    size: int = 16,
    color: RGBColor = RGBColor(0, 112, 192),
    orange_hex: str = "ED7D31",
    after_pt: float = 6,
) -> Paragraph:
    p = doc.add_paragraph()
    try:
        p.style = "Heading 1"
    except Exception:
        pass

    run = p.add_run(s(text))
    set_run(run, font, size, bold=True, color=color)
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1.0)
    _set_paragraph_bottom_border(p, color_hex=orange_hex, size_eighths=12, space=2)
    return p


def _emu_to_twips(emu: int) -> int:
    return int(round(int(emu) / 635.0))


def set_table_fixed_layout(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def set_table_width_exact(table, width) -> None:
    width_emu = int(width.emu)
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_emu_to_twips(width_emu)))

    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), "0")


def set_table_borders(table, *, color_hex: str = "A6A6A6") -> None:
    color_hex = s(color_hex).replace("#", "") or "A6A6A6"
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def set_table_columns_exact(table, widths_in: List[float]) -> None:
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    gridCols = tblGrid.findall(qn("w:gridCol"))

    while len(gridCols) < len(widths_in):
        gc = OxmlElement("w:gridCol")
        tblGrid.append(gc)
        gridCols = tblGrid.findall(qn("w:gridCol"))

    for i, w_in in enumerate(widths_in):
        w_emu = int(Inches(float(w_in)).emu)
        gridCols[i].set(qn("w:w"), str(_emu_to_twips(w_emu)))


def shade_cell(cell, fill_hex: str) -> None:
    fill_hex = s(fill_hex).replace("#", "") or "FFFFFF"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_margins(cell, top_dxa: int, bottom_dxa: int, left_dxa: int, right_dxa: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _set(side: str, val: int) -> None:
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")

    _set("top", int(top_dxa))
    _set("bottom", int(bottom_dxa))
    _set("start", int(left_dxa))
    _set("end", int(right_dxa))


def write_cell_text(
    cell,
    text: Any,
    *,
    font: str = "Times New Roman",
    size: int = 11,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1.0)
    r = p.add_run(s(text))
    set_run(r, font, size, bold=bold)


def write_yes_no_checkboxes(cell, value: Any, *, font_size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT, font: str = "Times New Roman") -> None:
    b = parse_bool_like(value)
    yes = (b is True)
    no = (b is False)

    cell.text = ""
    p = cell.paragraphs[0]
    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1.0)
    r = p.add_run(f"{_checkbox(yes)} Yes   {_checkbox(no)} No")
    set_run(r, font, font_size, bold=False)


def write_two_option_checkboxes(cell, value: Any, opt1: str, opt2: str, *, font_size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT, font: str = "Times New Roman") -> None:
    sv = s(value).lower().strip()
    o1 = opt1.lower().strip()
    o2 = opt2.lower().strip()
    c1 = bool(sv) and (sv == o1 or o1 in sv)
    c2 = bool(sv) and (sv == o2 or o2 in sv)

    cell.text = ""
    p = cell.paragraphs[0]
    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1.0)
    r = p.add_run(f"{_checkbox(c1)} {opt1}   {_checkbox(c2)} {opt2}")
    set_run(r, font, font_size, bold=False)


# ============================================================
# Data formatting helpers (isolated)
# ============================================================
def format_af_phone(v: Any) -> str:
    sv = s(v)
    if not sv:
        return ""
    digits = re.sub(r"\D+", "", sv)
    if len(digits) == 10 and digits.startswith("0"):
        return "+93" + digits[1:]
    if digits.startswith("93") and len(digits) in (11, 12):
        return "+" + digits
    if sv.startswith("+"):
        return sv
    return sv


def normalize_email_or_na_strict(v: Any) -> str:
    sv = s(v)
    if not sv:
        return "N/A"
    sv = sv.strip()
    if "@" not in sv or "." not in sv.split("@")[-1]:
        return "N/A"
    return sv


def donor_upper_and_pipe(v: Any) -> str:
    sv = s(v)
    if not sv:
        return ""
    sv = sv.replace("/", "|").replace(",", "|").replace(";", "|")
    parts = [p.strip() for p in sv.split("|") if p.strip()]
    return " | ".join([p.upper() for p in parts])


def format_date_dd_mon_yyyy(value: Any) -> str:
    sv = s(value)
    if not sv:
        return ""
    sv_clean = sv.split(".")[0].replace("T", " ").replace("Z", "").strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            return datetime.strptime(sv_clean, fmt).strftime("%d/%b/%Y")
        except Exception:
            pass
    return sv.split(" ")[0] if " " in sv else sv


# ============================================================
# Config (easy to extend)
# ============================================================
TITLE_TEXT = "1.  General Project Information:"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16
TITLE_BLUE = RGBColor(0, 112, 192)

BODY_FONT = "Times New Roman"
BODY_SIZE = 11

FIELD_FILL_HEX = "D9E2F3"
HEADER_FILL_HEX = "E7EEF9"
BORDER_HEX = "A6A6A6"

FIELD_COL_WIDTH_IN = 2.30
M_TOP, M_BOTTOM, M_LEFT, M_RIGHT = 80, 80, 120, 120


# Google-sheet keys used by this section (only the ones that matter here)
DATA_KEYS: Dict[str, str] = {
    "PROVINCE": "A01_Province",
    "DISTRICT": "A02_District",
    "VILLAGE": "Village",
    "GPS_LAT": "GPS_1-Latitude",
    "GPS_LON": "GPS_1-Longitude",
    "STARTTIME": "starttime",
    "ACTIVITY_NAME": "Activity_Name",
    "PRIMARY_PARTNER": "Primary_Partner_Name",
    "MONITOR_NAME": "A07_Monitor_name",
    "MONITOR_EMAIL": "A12_Monitor_email",
    "RESPONDENT_NAME": "A08_Respondent_name",
    "RESPONDENT_PHONE": "A10_Respondent_phone",
    "RESPONDENT_EMAIL": "A11_Respondent_email",
    "RESPONDENT_SEX": "A09_Respondent_sex",
    "PROJECT_COST_LABEL": "A14_Estimated_cost_amount_label",
    "EST_COST_AMOUNT": "Estimated_Project_Cost_amount",
    "CONTRACT_COST_AMOUNT": "Contracted_Project_Cost_amount",
    "PROJECT_STATUS_LABEL": "Project_Status",
    "PROJECT_PROGRESS_LABEL": "Project_progress",
    "START_DATE": "A15_Contract_start_date",
    "END_DATE": "A16_Contract_end_date",
    "PREV_PROGRESS": "A17_Previous_physical_progress",
    "CURR_PROGRESS": "A18_Current_physical_progress",
    "DONOR_NAME": "A24_Donor_name",
    "MONITORING_REPORT_NO": "A25_Monitoring_report_number",
    "CURRENT_REPORT_DATE": "A20_Current_report_date",
    "PREV_REPORT_DATE": "A21_Last_report_date",
    "VISIT_NO": "A26_Visit_number",
    "DOC_CONTRACT": "B3_Contract",
    "DOC_JOURNAL": "B4_Journal",
    "DOC_BOQ": "D2_boq_available",
    "DOC_DRAWINGS": "B1_Design_drawings",
    "DOC_SITE_ENGINEER": "B6_Site_engineer",
    "DOC_GEOPHYSICAL": "D3_geophysical_tests_available",
    "DOC_WQ_TEST": "D4_water_quality_tests_available",
    "DOC_PUMP_TEST": "D4_pump_test_results_available",
    "COMMUNITY_AGREEMENT": "community_agreement",
    "WORK_SAFETY": "work_safety_considered",
    "ENV_RISK": "environmental_risk",
}

# IMPORTANT: override keys MUST match Step2 field labels exactly.
OVK = {
    "Province": "Province",
    "District": "District",
    "Village": "Village / Community",
    "GPS": "GPS points",
    "Project Name": "Project Name",
    "Date of Visit": "Date of Visit",
    "IP": "Name of the IP, Organization / NGO",
    "Monitor Name": "Name of the monitor engineer",   # Step2 key
    "Monitor Email": "Email of the monitor engineer",
    "Respondent Name": "Name of the respondent (Participant / UNICEF / IPs)",
    "Respondent Phone": "Contact Number of the Respondent",
    "Respondent Email": "Email Address of the Respondent",
    "Respondent Sex": "Sex of Respondent",
    "Estimated Cost": "Estimated Project Cost",
    "Contracted Cost": "Contracted Project Cost",
    "Project Status": "Project Status",
    "Reason Delay": "Reason for delay",
    "Project Progress": "Project progress",
    "Contract Start": "Contract Start Date",
    "Contract End": "Contract End Date",
    "Prev Prog": "Previous Physical Progress (%)",
    "Curr Prog": "Current Physical Progress (%)",
    "CDC": "CDC Code",
    "Donor": "Donor Name",
    "Report No": "Monitoring Report Number",
    "Current Report Date": "Date of Current Report",
    "Last Report Date": "Date of Last Monitoring Report",
    "Sites": "Number of Sites Visited",
    # Docs are stored as Yes/No per doc field name in Step2 (Contract, Journal, BOQ, Design drawings, ...)
}

# Documents: label in report -> sheet key fallback
DOCS: List[Tuple[str, str]] = [
    ("Contract", DATA_KEYS["DOC_CONTRACT"]),
    ("Journal", DATA_KEYS["DOC_JOURNAL"]),
    ("BOQ", DATA_KEYS["DOC_BOQ"]),
    ("Design drawings", DATA_KEYS["DOC_DRAWINGS"]),
    ("Site engineer", DATA_KEYS["DOC_SITE_ENGINEER"]),
    ("Geophysical tests", DATA_KEYS["DOC_GEOPHYSICAL"]),
    ("Water quality tests", DATA_KEYS["DOC_WQ_TEST"]),
    ("Pump test results", DATA_KEYS["DOC_PUMP_TEST"]),
]


@dataclass(frozen=True)
class FieldSpec:
    label: str
    override_key: Optional[str] = None
    sheet_key: Optional[str] = None
    formatter: Optional[Callable[[Any], str]] = None


def _pick(overrides: Dict[str, Any], row: Dict[str, Any], *, override_key: Optional[str], sheet_key: Optional[str], fallback: Any = "") -> Any:
    if override_key:
        v = overrides.get(override_key, None)
        if v is not None and s(v) != "":
            return v
    if sheet_key:
        v2 = row.get(sheet_key, None)
        if v2 is not None and s(v2) != "":
            return v2
    return fallback


def _set_a4_narrow(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)


def _usable_width_inches(section) -> float:
    usable_emu = section.page_width.emu - section.left_margin.emu - section.right_margin.emu
    return float(usable_emu) / 914400.0


# ============================================================
# Available documents inner table (borderless)
# ============================================================
def add_available_documents_inner_table(cell, *, row: Dict[str, Any], overrides: Dict[str, Any]) -> None:
    cell.text = ""

    docs: List[Tuple[str, Any]] = []
    for label, sheet_key in DOCS:
        # Step2 stores doc answers by label exactly (e.g., "Contract": "Yes"/"No")
        v = overrides.get(label, None)
        if v is None or s(v) == "":
            v = row.get(sheet_key, None)
        docs.append((label, v))

    inner = cell.add_table(rows=len(docs), cols=2)
    inner.autofit = False
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_fixed_layout(inner)

    # remove borders
    tbl = inner._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")

    try:
        set_table_columns_exact(inner, [3.2, 1.2])
    except Exception:
        pass

    for i, (label, val) in enumerate(docs):
        r = inner.rows[i]
        c0, c1 = r.cells[0], r.cells[1]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_margins(c0, 40, 40, 80, 80)
        set_cell_margins(c1, 40, 40, 80, 80)

        write_cell_text(c0, label, font=BODY_FONT, size=BODY_SIZE, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

        available = _truthy_doc(val)
        c1.text = ""
        p = c1.paragraphs[0]
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1.0)
        rr = p.add_run(f"{_checkbox(available)} Yes   {_checkbox(not available)} No")
        set_run(rr, BODY_FONT, BODY_SIZE, bold=False)


# ============================================================
# Main section
# ============================================================
def add_general_project_information(
    doc: Document,
    row: Dict[str, Any],
    overrides: Optional[Dict[str, Any]] = None,
    respondent_sex_val: Any = None,
) -> None:
    overrides = overrides or {}
    row = row or {}

    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    _set_a4_narrow(section)

    add_section_title_h1(
        doc,
        TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=TITLE_BLUE,
        orange_hex="ED7D31",
        after_pt=6,
    )

    # ---- Build values (override first, then sheet) ----
    gps_lat = s(row.get(DATA_KEYS["GPS_LAT"]))
    gps_lon = s(row.get(DATA_KEYS["GPS_LON"]))

    # Costs: keep your legacy logic
    cost_label = s(row.get(DATA_KEYS["PROJECT_COST_LABEL"])).lower()
    estimated_amount = s(row.get(DATA_KEYS["EST_COST_AMOUNT"]))
    contracted_amount = s(row.get(DATA_KEYS["CONTRACT_COST_AMOUNT"]))
    estimated_cost_sheet = estimated_amount if "estimated" in cost_label else ""
    contracted_cost_sheet = contracted_amount if "contract" in cost_label else ""

    # community/work/env keys: keep legacy fallbacks too
    community_sheet = row.get(DATA_KEYS["COMMUNITY_AGREEMENT"]) or row.get("Community_agreement")
    work_sheet = row.get(DATA_KEYS["WORK_SAFETY"])
    env_sheet = row.get(DATA_KEYS["ENV_RISK"])

    specs: List[FieldSpec] = [
        FieldSpec("Province", override_key=OVK["Province"], sheet_key=DATA_KEYS["PROVINCE"]),
        FieldSpec("District", override_key=OVK["District"], sheet_key=DATA_KEYS["DISTRICT"]),
        FieldSpec("Village / Community", override_key=OVK["Village"], sheet_key=DATA_KEYS["VILLAGE"]),
        FieldSpec("GPS points", override_key=OVK["GPS"], sheet_key=None, formatter=lambda _: f"{gps_lat}, {gps_lon}".strip().strip(",")),
        FieldSpec("Project Name", override_key=OVK["Project Name"], sheet_key=DATA_KEYS["ACTIVITY_NAME"]),
        FieldSpec("Date of Visit", override_key=OVK["Date of Visit"], sheet_key=None, formatter=lambda _: s(overrides.get(OVK["Date of Visit"], format_date_dd_mon_yyyy(row.get(DATA_KEYS["STARTTIME"])) ))),
        FieldSpec("Name of the IP, Organization / NGO", override_key=OVK["IP"], sheet_key=DATA_KEYS["PRIMARY_PARTNER"]),
        FieldSpec("Name of the monitor Engineer", override_key=OVK["Monitor Name"], sheet_key=DATA_KEYS["MONITOR_NAME"]),
        FieldSpec("Email of the monitor engineer", override_key=OVK["Monitor Email"], sheet_key=DATA_KEYS["MONITOR_EMAIL"]),
        FieldSpec("Name of the respondent (Participant / UNICEF / IPs)", override_key=OVK["Respondent Name"], sheet_key=DATA_KEYS["RESPONDENT_NAME"]),
        FieldSpec("Contact Number of the Respondent", override_key=OVK["Respondent Phone"], sheet_key=None, formatter=lambda _: s(overrides.get(OVK["Respondent Phone"], format_af_phone(row.get(DATA_KEYS["RESPONDENT_PHONE"])) ))),
        FieldSpec("Email Address of the Respondent", override_key=OVK["Respondent Email"], sheet_key=None, formatter=lambda _: s(overrides.get(OVK["Respondent Email"], normalize_email_or_na_strict(row.get(DATA_KEYS["RESPONDENT_EMAIL"])) ))),
        FieldSpec("Estimated Project Cost", override_key=OVK["Estimated Cost"], sheet_key=None, formatter=lambda _: na(overrides.get(OVK["Estimated Cost"], estimated_cost_sheet))),
        FieldSpec("Contracted Project Cost", override_key=OVK["Contracted Cost"], sheet_key=None, formatter=lambda _: na(overrides.get(OVK["Contracted Cost"], contracted_cost_sheet))),
        FieldSpec("Project Status", override_key=OVK["Project Status"], sheet_key=DATA_KEYS["PROJECT_STATUS_LABEL"], formatter=lambda v: three_option_checkbox_line(v, "Ongoing", "Completed", "Suspended")),
        FieldSpec("Reason for delay", override_key=OVK["Reason Delay"], sheet_key="B8_Reasons_for_delay", formatter=lambda v: na(v)),
        FieldSpec("Project progress", override_key=OVK["Project Progress"], sheet_key=DATA_KEYS["PROJECT_PROGRESS_LABEL"], formatter=lambda v: three_option_checkbox_line(v, "Ahead of Schedule", "On Schedule", "Running behind")),
        FieldSpec("Contract Start Date", override_key=OVK["Contract Start"], sheet_key=None, formatter=lambda _: na(overrides.get(OVK["Contract Start"], format_date_dd_mon_yyyy(row.get(DATA_KEYS["START_DATE"])) ))),
        FieldSpec("Contract End Date", override_key=OVK["Contract End"], sheet_key=None, formatter=lambda _: na(overrides.get(OVK["Contract End"], format_date_dd_mon_yyyy(row.get(DATA_KEYS["END_DATE"])) ))),
        FieldSpec("Previous Physical Progress (%)", override_key=OVK["Prev Prog"], sheet_key=DATA_KEYS["PREV_PROGRESS"], formatter=lambda v: na(v)),
        FieldSpec("Current Physical Progress (%)", override_key=OVK["Curr Prog"], sheet_key=DATA_KEYS["CURR_PROGRESS"], formatter=lambda v: na(v)),
        FieldSpec("CDC Code", override_key=OVK["CDC"], sheet_key="A23_CDC_code"),
        FieldSpec("Donor Name", override_key=OVK["Donor"], sheet_key=DATA_KEYS["DONOR_NAME"], formatter=lambda v: donor_upper_and_pipe(v)),
        FieldSpec("Monitoring Report Number", override_key=OVK["Report No"], sheet_key=DATA_KEYS["MONITORING_REPORT_NO"]),
        FieldSpec("Date of Current Report", override_key=OVK["Current Report Date"], sheet_key=None, formatter=lambda _: s(overrides.get(OVK["Current Report Date"], format_date_dd_mon_yyyy(row.get(DATA_KEYS["CURRENT_REPORT_DATE"])) ))),
        FieldSpec("Date of Last Monitoring Report", override_key=OVK["Last Report Date"], sheet_key=None, formatter=lambda _: na(overrides.get(OVK["Last Report Date"], format_date_dd_mon_yyyy(row.get(DATA_KEYS["PREV_REPORT_DATE"])) ))),
        FieldSpec("Number of Sites Visited", override_key=OVK["Sites"], sheet_key=DATA_KEYS["VISIT_NO"]),
    ]

    # respondent sex special (checkbox rendering)
    if respondent_sex_val is None:
        respondent_sex_val = overrides.get(OVK["Respondent Sex"], row.get(DATA_KEYS["RESPONDENT_SEX"]))

    # ---- Table ----
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    set_table_fixed_layout(table)
    set_table_borders(table, color_hex=BORDER_HEX)

    usable_w_in = _usable_width_inches(section)
    set_table_width_exact(table, Inches(usable_w_in))

    field_in = min(float(FIELD_COL_WIDTH_IN), max(1.0, usable_w_in - 1.0))
    detail_in = max(0.75, usable_w_in - field_in)
    set_table_columns_exact(table, [field_in, detail_in])

    field_w = Inches(field_in)
    detail_w = Inches(detail_in)

    hdr = table.rows[0]
    for i, txt in enumerate(("Field", "Details")):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, HEADER_FILL_HEX)
        set_cell_margins(cell, M_TOP, M_BOTTOM, M_LEFT, M_RIGHT)
        write_cell_text(cell, txt, font=BODY_FONT, size=BODY_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    def add_row(field_label: str, value_text: str) -> None:
        r = table.add_row()
        c0, c1 = r.cells[0], r.cells[1]
        c0.width = field_w
        c1.width = detail_w

        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c0, FIELD_FILL_HEX)
        set_cell_margins(c0, M_TOP, M_BOTTOM, M_LEFT, M_RIGHT)
        write_cell_text(c0, field_label, font=BODY_FONT, size=BODY_SIZE, bold=True)

        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(c1, M_TOP, M_BOTTOM, M_LEFT, M_RIGHT)
        write_cell_text(c1, value_text, font=BODY_FONT, size=BODY_SIZE, bold=False)

        set_table_columns_exact(table, [field_in, detail_in])  # keep layout stable

    def add_row_custom(field_label: str, renderer) -> None:
        r = table.add_row()
        c0, c1 = r.cells[0], r.cells[1]
        c0.width = field_w
        c1.width = detail_w

        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c0, FIELD_FILL_HEX)
        set_cell_margins(c0, M_TOP, M_BOTTOM, M_LEFT, M_RIGHT)
        write_cell_text(c0, field_label, font=BODY_FONT, size=BODY_SIZE, bold=True)

        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(c1, M_TOP, M_BOTTOM, M_LEFT, M_RIGHT)
        c1.text = ""
        renderer(c1)

        set_table_columns_exact(table, [field_in, detail_in])

    # write normal rows from specs
    for spec in specs:
        if spec.formatter is not None:
            v = spec.formatter(_pick(overrides, row, override_key=spec.override_key, sheet_key=spec.sheet_key))
        else:
            v = s(_pick(overrides, row, override_key=spec.override_key, sheet_key=spec.sheet_key))
        add_row(spec.label, s(v))

    # Sex checkbox row
    add_row_custom(
        "Sex of Respondent",
        lambda cell: write_two_option_checkboxes(cell, respondent_sex_val, "Male", "Female", font_size=BODY_SIZE),
    )

    # Available documents
    add_row_custom("Available documents in the site", lambda cell: add_available_documents_inner_table(cell, row=row, overrides=overrides))

    # Yes/No checkbox rows
    community_val = overrides.get(
        "Community agreement (Is the community/user group agreed on the well site?)",
        community_sheet,
    )
    work_val = overrides.get("Work safety considered", work_sheet)
    env_val = overrides.get("Environmental risk", env_sheet)

    add_row_custom(
        "Community agreement (Is the community/user group agreed on the well site?)",
        lambda cell: write_yes_no_checkboxes(cell, community_val, font_size=BODY_SIZE),
    )
    add_row_custom(
        "Work safety considered",
        lambda cell: write_yes_no_checkboxes(cell, work_val, font_size=BODY_SIZE),
    )
    add_row_custom(
        "Environmental risk",
        lambda cell: write_yes_no_checkboxes(cell, env_val, font_size=BODY_SIZE),
    )
