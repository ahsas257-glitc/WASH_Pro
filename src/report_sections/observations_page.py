from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional, Tuple

from PIL import Image
from docx.document import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# =============================================================================
# CONSTANTS (layout)
# =============================================================================
# Photo box in Word (exact requirement)
PHOTO_W_IN = 3.50
PHOTO_ASPECT = 3.17 / 2.38
PHOTO_H_IN = PHOTO_W_IN / PHOTO_ASPECT

# Two-column block for each photo: LEFT text + RIGHT image
LEFT_TEXT_W_IN = 3.64
RIGHT_PHOTO_W_IN = 3.64

# Major findings table widths
MF_COL_NO = 0.38
MF_COL_FIND = 2.88
MF_COL_COMP = 0.91
MF_COL_PHOTO = 3.13
MF_COL_WIDTHS = [MF_COL_NO, MF_COL_FIND, MF_COL_COMP, MF_COL_PHOTO]

MF_HEADER_FILL_HEX = "2F5597"
MF_BORDER_HEX = "000000"


# =============================================================================
# BASIC helpers
# =============================================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def _compact(p) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def _one_line_gap(doc: Document) -> None:
    p = doc.add_paragraph("")
    _compact(p)


def _ensure_paragraph_font_size(p, size_pt: int) -> None:
    """
    Force font size on runs in a paragraph (helps keep consistent output).
    Keeps heading COLOR intact (we do not touch run color).
    """
    try:
        for run in (p.runs or []):
            run.font.size = Pt(float(size_pt))
    except Exception:
        pass


# =============================================================================
# Table helpers (stable widths + borders)
# =============================================================================
def _set_table_fixed_layout(tbl) -> None:
    try:
        tbl.allow_autofit = False  # type: ignore[attr-defined]
    except Exception:
        pass
    tbl.autofit = False

    tblPr = tbl._tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _set_tbl_width_in(tbl, width_in: float) -> None:
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(float(width_in) * 1440)))  # inches -> twips


def _set_cell_width_in(cell, width_in: float) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(float(width_in) * 1440)))


def _apply_2col_widths_hard(tbl, left_w_in: float, right_w_in: float) -> None:
    """
    Hard lock a 2-column table so text doesn't get squeezed by pictures.
    """
    _set_table_fixed_layout(tbl)
    _set_tbl_width_in(tbl, float(left_w_in) + float(right_w_in))

    row = tbl.rows[0]
    _set_cell_width_in(row.cells[0], left_w_in)
    _set_cell_width_in(row.cells[1], right_w_in)

    # python-docx widths (best effort)
    try:
        tbl.columns[0].width = Inches(float(left_w_in))
        tbl.columns[1].width = Inches(float(right_w_in))
        row.cells[0].width = Inches(float(left_w_in))
        row.cells[1].width = Inches(float(right_w_in))
    except Exception:
        pass


def _remove_table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for k in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = borders.find(qn(f"w:{k}"))
        if e is None:
            e = OxmlElement(f"w:{k}")
            borders.append(e)
        e.set(qn("w:val"), "nil")


def _apply_table_col_widths(tbl, widths_in: List[float]) -> None:
    """
    Lock columns widths (fast, stable). For each column, also sets each cell width.
    """
    _set_table_fixed_layout(tbl)
    for i, w in enumerate(widths_in):
        try:
            tbl.columns[i].width = Inches(float(w))
            for cell in tbl.columns[i].cells:
                cell.width = Inches(float(w))
        except Exception:
            pass


def _set_table_borders(tbl, *, color_hex: str = MF_BORDER_HEX, size: str = "12") -> None:
    color_hex = s(color_hex).replace("#", "") or "000000"
    t = tbl._tbl
    tbl_pr = t.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def _shade_cell(cell, fill_hex: str) -> None:
    fill_hex = s(fill_hex).replace("#", "") or "FFFFFF"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_margins(cell, start: int = 70, end: int = 70, top: int = 0, bottom: int = 0) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, val in (("start", start), ("end", end), ("top", top), ("bottom", bottom)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(int(val)))
        node.set(qn("w:type"), "dxa")


# =============================================================================
# Image helpers (safe + predictable)
# =============================================================================
def _bytes_look_like_html(data: bytes) -> bool:
    head = (data or b"")[:300].lower()
    return b"<!doctype html" in head or b"<html" in head or b"<head" in head


def _crop_to_aspect(img: Image.Image, target_aspect: float) -> Image.Image:
    w, h = img.size
    if w <= 0 or h <= 0:
        return img

    current = w / h
    if abs(current - target_aspect) < 1e-3:
        return img

    if current > target_aspect:
        # crop width
        new_w = int(h * target_aspect)
        left = max(0, (w - new_w) // 2)
        return img.crop((left, 0, left + new_w, h))

    # crop height
    new_h = int(w * (1 / target_aspect))
    top = max(0, (h - new_h) // 2)
    return img.crop((0, top, w, top + new_h))


def _to_clean_png_fit_box(b: Optional[bytes], *, target_aspect: float) -> Optional[bytes]:
    """
    Converts bytes -> cropped PNG (fast enough; no resizing here to preserve clarity).
    Returns None if bytes invalid/unreadable.
    """
    if not b or _bytes_look_like_html(b):
        return None
    try:
        img = Image.open(io.BytesIO(b)).convert("RGB")
        img = _crop_to_aspect(img, target_aspect)

        out = io.BytesIO()
        img.save(out, format="PNG", optimize=True)
        return out.getvalue()
    except Exception:
        return None


# =============================================================================
# Text helpers
# =============================================================================
def _extract_obs_no_prefix(obs_title: str) -> str:
    """
    Extract "5.1" from "5.1. Title ...".
    """
    t = s(obs_title)
    m = re.match(r"^\s*(\d+\.\d+)\.", t)
    return m.group(1) if m else ""


def _write_cell_text(
    cell,
    text: Any,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    bold: bool = False,
    italic: bool = False,
    size_pt: int = 10,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    _compact(p)

    r = p.add_run(s(text))
    r.bold = bool(bold)
    r.italic = bool(italic)
    try:
        r.font.size = Pt(float(size_pt))
    except Exception:
        pass


# =============================================================================
# Major findings table
# =============================================================================
def _add_major_findings_table(
    doc: Document,
    *,
    major_rows: List[Dict[str, Any]],
    photo_bytes: Dict[str, bytes],
) -> None:
    """
    major_rows item shape:
      {"finding": "...", "compliance": "...", "photo": "url", "photo_bytes": b"..."}  # photo_bytes optional
    """
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    _set_table_fixed_layout(tbl)
    _set_table_borders(tbl, color_hex=MF_BORDER_HEX, size="12")
    _apply_table_col_widths(tbl, MF_COL_WIDTHS)

    headers = ["NO", "Findings", "Compliance", "Photos"]
    hdr_cells = tbl.rows[0].cells
    for i, txt in enumerate(headers):
        c = hdr_cells[i]
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(c, MF_HEADER_FILL_HEX)
        _set_cell_margins(c, start=90, end=90, top=50, bottom=50)

        p = c.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _compact(p)

        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        # keep the white header look (consistent)
        try:
            from docx.shared import RGBColor

            run.font.color.rgb = RGBColor(255, 255, 255)
        except Exception:
            pass

    for idx, rdata in enumerate(major_rows or [], start=1):
        if not isinstance(rdata, dict):
            continue

        row = tbl.add_row().cells
        for c in row:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(c, start=90, end=90, top=50, bottom=50)

        _write_cell_text(row[0], str(idx), align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=9)
        _write_cell_text(row[1], rdata.get("finding"), align=WD_ALIGN_PARAGRAPH.JUSTIFY, size_pt=9)
        _write_cell_text(row[2], rdata.get("compliance"), align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=9)

        # photo cell (right aligned)
        row[3].text = ""
        pph = row[3].paragraphs[0]
        pph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _compact(pph)

        url = s(rdata.get("photo"))
        embedded = rdata.get("photo_bytes")
        if isinstance(embedded, bytearray):
            embedded = bytes(embedded)

        b = embedded if isinstance(embedded, (bytes, bytearray)) and embedded else (photo_bytes.get(url) if url else None)
        clean = _to_clean_png_fit_box(b, target_aspect=PHOTO_ASPECT) if b else None
        if clean:
            pph.add_run().add_picture(io.BytesIO(clean), width=Inches(MF_COL_PHOTO - 0.05))


# =============================================================================
# PUBLIC API
# =============================================================================
def add_observations_page(
    doc: Document,
    *,
    component_observations: List[Dict[str, Any]],
    photo_bytes: Dict[str, bytes],
) -> None:
    """
    Input must be what your Step3 stores:
      component_observations = [
        {
          "comp_id": "...",
          "title": "...",
          "observations_valid": [
            {
              "title": "5.1. ...",
              "audio_url": "...",
              "photos": [{"url": "...", "text": "...", "bytes": b"..."}],
              "major_table": [...],
              "recommendations": [...],
            }
          ]
        }
      ]

    Guarantee:
      - User edits in Step3 (photo text) are already in obs["photos"][i]["text"] => printed.
      - If bytes embedded in obs photos, uses them; else uses `photo_bytes[url]` cache.
      - Heading colors remain "natural" because we DO NOT override run color for headings.
    """
    if not component_observations:
        return

    photo_bytes = photo_bytes or {}

    # New page for section 5
    doc.add_page_break()

    # Heading 1: keep natural Heading 1 styling/color (do NOT set RGBColor here)
    h = doc.add_paragraph("5. Project Component Wise Key Observations:")
    try:
        h.style = "Heading 1"
    except Exception:
        pass
    _compact(h)
    _ensure_paragraph_font_size(h, 16)

    for comp in component_observations or []:
        if not isinstance(comp, dict):
            continue

        comp_id = s(comp.get("comp_id"))
        comp_title = s(comp.get("title"))
        comp_line = f"{comp_id} — {comp_title}".strip(" —")

        ch = doc.add_paragraph(comp_line)
        try:
            ch.style = "Heading 2"
        except Exception:
            pass
        _compact(ch)
        _ensure_paragraph_font_size(ch, 14)

        for obs in comp.get("observations_valid", []) or []:
            if not isinstance(obs, dict):
                continue

            # Observation title (Heading 2, keep natural color)
            obs_title = s(obs.get("title"))
            ot = doc.add_paragraph(obs_title)
            try:
                ot.style = "Heading 2"
            except Exception:
                pass
            _compact(ot)
            _ensure_paragraph_font_size(ot, 14)

            photos = obs.get("photos") or []
            if not isinstance(photos, list):
                photos = []

            if not photos:
                p = doc.add_paragraph("No photos were selected for this observation.")
                _compact(p)
                _one_line_gap(doc)
            else:
                for i, ph in enumerate(photos):
                    if not isinstance(ph, dict):
                        continue

                    url = s(ph.get("url"))
                    txt = s(ph.get("text"))

                    embedded = ph.get("bytes")
                    if isinstance(embedded, bytearray):
                        embedded = bytes(embedded)

                    b = embedded if isinstance(embedded, (bytes, bytearray)) and embedded else (photo_bytes.get(url) if url else None)
                    clean = _to_clean_png_fit_box(b, target_aspect=PHOTO_ASPECT) if b else None

                    # 2-col table per photo
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    _remove_table_borders(tbl)

                    # hard widths so left text stays visible
                    _apply_2col_widths_hard(tbl, LEFT_TEXT_W_IN, RIGHT_PHOTO_W_IN)

                    left, right = tbl.rows[0].cells
                    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                    _set_cell_margins(left, start=80, end=120)
                    _set_cell_margins(right, start=60, end=60)

                    # LEFT: user-entered observation text (exactly what they typed)
                    _write_cell_text(
                        left,
                        txt,
                        align=WD_ALIGN_PARAGRAPH.LEFT,
                        bold=False,
                        italic=False,
                        size_pt=10,
                    )

                    # RIGHT: photo
                    right.paragraphs[0].text = ""
                    _compact(right.paragraphs[0])

                    if not url:
                        p = right.add_paragraph("Photo missing: empty url")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _compact(p)
                    elif not b:
                        p = right.add_paragraph(f"Photo missing (not cached): {url}")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _compact(p)
                    elif not clean:
                        p = right.add_paragraph(f"Photo unreadable: {url}")
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _compact(p)
                    else:
                        pic_p = right.add_paragraph()
                        pic_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _compact(pic_p)
                        pic_p.add_run().add_picture(
                            io.BytesIO(clean),
                            width=Inches(PHOTO_W_IN),
                            height=Inches(PHOTO_H_IN),
                        )

                    if i < len(photos) - 1:
                        _one_line_gap(doc)

                _one_line_gap(doc)

            # Major findings + Recommendations (optional)
            base_no = _extract_obs_no_prefix(obs_title)
            major_title = f"{base_no}.1 Major findings:" if base_no else "Major findings:"
            reco_title = f"{base_no}.2 Recommendations:" if base_no else "Recommendations:"

            major_rows = obs.get("major_table") or []
            if isinstance(major_rows, list) and len(major_rows) > 0:
                mh = doc.add_paragraph(major_title)
                try:
                    mh.style = "Heading 3"
                except Exception:
                    pass
                _compact(mh)
                _ensure_paragraph_font_size(mh, 12)

                _one_line_gap(doc)
                _add_major_findings_table(doc, major_rows=major_rows, photo_bytes=photo_bytes)
                _one_line_gap(doc)

            recs = obs.get("recommendations") or []
            if isinstance(recs, list) and any(s(x) for x in recs):
                rh = doc.add_paragraph(reco_title)
                try:
                    rh.style = "Heading 3"
                except Exception:
                    pass
                _compact(rh)
                _ensure_paragraph_font_size(rh, 12)

                _one_line_gap(doc)
                for t in recs:
                    tx = s(t)
                    if not tx:
                        continue
                    rp = doc.add_paragraph(tx, style="List Bullet")
                    _compact(rp)
                _one_line_gap(doc)
