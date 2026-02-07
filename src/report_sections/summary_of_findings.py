from __future__ import annotations

from typing import Any, Dict, List, Optional, Union

from docx.document import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor, Pt
from docx.text.paragraph import Paragraph


# ============================================================
# Inline helpers
# ============================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: Union[int, float] = 0,
    after_pt: Union[int, float] = 0,
    line_spacing: float = 1.0,
) -> None:
    try:
        p.alignment = align
    except Exception:
        pass
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def set_run(
    run,
    font: str,
    size: Union[int, float],
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    try:
        run.font.name = font
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass
    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass
    try:
        run.bold = bool(bold)
    except Exception:
        pass
    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


# ============================================================
# Border thickness (Word uses eighths of a point)
# 1/4 pt = 0.25 pt => 0.25 * 8 = 2
# ============================================================
BORDER_SZ_EIGHTHS = 2  # ✅ 1/4pt


# -------------------------
# Heading 1 with orange underline (TOC-friendly)
# -------------------------
def _set_paragraph_bottom_border(
    paragraph: Paragraph,
    *,
    color_hex: str,
    size_eighths: int = 12,
    space: int = 2,
) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), s(color_hex).replace("#", "") or "000000")


def add_section_title_h1(
    doc: Document,
    text: str,
    *,
    font: str = "Cambria",
    size: int = 16,
    color: RGBColor = RGBColor(0, 112, 192),
    orange_hex: str = "ED7D31",
    after_pt: Union[int, float] = 6,
) -> Paragraph:
    p = doc.add_paragraph()
    try:
        p.style = "Heading 1"
    except Exception:
        pass

    r = p.add_run(s(text))
    set_run(r, font, size, bold=True, color=color)
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1.0)
    _set_paragraph_bottom_border(p, color_hex=orange_hex, size_eighths=12, space=2)
    return p


# -------------------------
# Table/layout helpers
# -------------------------
def _emu_to_twips(emu: int) -> int:
    return int(round(int(emu) / 635.0))


def section_usable_width_emu(section) -> int:
    try:
        return int(section.page_width.emu - section.left_margin.emu - section.right_margin.emu)
    except Exception:
        try:
            return int(section.page_width - section.left_margin - section.right_margin)
        except Exception:
            return 914400 * 6


def _safe_section(doc: Document, idx: int = 0):
    try:
        return doc.sections[idx]
    except Exception:
        return doc.sections[0]


def set_table_fixed_layout(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)

    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def set_table_width_from_section(table, doc: Document, *, section_index: int = 0) -> None:
    section = _safe_section(doc, section_index)
    usable_emu = section_usable_width_emu(section)

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_emu_to_twips(usable_emu)))

    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), "0")


def set_table_borders(table, *, color_hex: str = "000000", size_eighths: int = BORDER_SZ_EIGHTHS) -> None:
    """
    Word border size is in eighths of a point (w:sz).
    """
    color_hex = s(color_hex).replace("#", "") or "000000"
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.append(tbl_pr)

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(int(size_eighths)))  # ✅ 1/4pt => 2
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def set_cell_borders(cell, *, size_eighths: int = BORDER_SZ_EIGHTHS, color_hex: str = "000000") -> None:
    """
    Cell borders in eighths of a point (w:sz).
    """
    color_hex = s(color_hex).replace("#", "") or "000000"
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(int(size_eighths)))  # ✅ 1/4pt => 2
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def shade_cell(cell, fill_hex: str) -> None:
    fill_hex = s(fill_hex).replace("#", "") or "FFFFFF"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_margins(
    cell,
    *,
    top_dxa: int = 80,
    bottom_dxa: int = 80,
    left_dxa: int = 120,
    right_dxa: int = 120,
) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _set(side: str, val: int) -> None:
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")

    _set("top", top_dxa)
    _set("bottom", bottom_dxa)
    _set("start", left_dxa)
    _set("end", right_dxa)


def set_row_cant_split(row, *, cant_split: bool = True) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    existing = trPr.find(qn("w:cantSplit"))
    if cant_split:
        if existing is None:
            trPr.append(OxmlElement("w:cantSplit"))
    else:
        if existing is not None:
            trPr.remove(existing)


def set_row_height_exact(row, height_in: float) -> None:
    row.height = Inches(float(height_in))
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def set_row_height_at_least(row, height_in: float) -> None:
    row.height = Inches(float(height_in))
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _scale_widths_to_fit(usable_in: float, desired: List[float]) -> List[float]:
    usable_in = max(5.5, float(usable_in or 0))
    widths = [max(0.05, float(x)) for x in desired]
    total = sum(widths)
    if total <= usable_in:
        return widths
    factor = usable_in / max(total, 1e-9)
    return [w * factor for w in widths]


# ============================================================
# Section 6 Style
# ============================================================
TITLE_TEXT = "6.        Summary of the findings:"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"

DARK_BLUE_HEX = "1F4E79"
BORDER_HEX = "000000"
HEADER_TEXT_COLOR = RGBColor(255, 255, 255)

FONT_BODY = "Times New Roman"
FONT_SIZE_BODY = 10
FONT_SIZE_HEADER = 10


def normalize_sentence(v: Any) -> str:
    sv = s(v)
    if not sv:
        return ""
    sv = " ".join(sv.split()).strip()
    if sv and sv not in ("—", "-"):
        if sv[-1] not in ".!?":
            sv += "."
    if sv and sv[0].isalpha():
        sv = sv[0].upper() + sv[1:]
    return sv


def _resolve_severity(
    idx: int,
    finding_text: str,
    severity_by_no: Dict[int, str],
    severity_by_finding: Dict[str, str],
) -> str:
    chosen = s(severity_by_no.get(idx)).strip()
    if chosen:
        return chosen.title()

    f_norm = normalize_sentence(finding_text).lower()
    for k, v in (severity_by_finding or {}).items():
        if normalize_sentence(k).lower() == f_norm:
            return s(v).title() or "Medium"
    return "Medium"


def add_summary_of_findings_section6(
    doc: Document,
    *,
    extracted_rows: Optional[List[Dict[str, str]]] = None,
    severity_by_no: Optional[Dict[int, str]] = None,
    severity_by_finding: Optional[Dict[str, str]] = None,
    add_legend: bool = True,
    add_page_break_before: bool = True,
) -> None:
    severity_by_no = severity_by_no or {}
    severity_by_finding = severity_by_finding or {}

    if add_page_break_before:
        doc.add_page_break()

    add_section_title_h1(
        doc,
        TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=TITLE_BLUE,
        orange_hex=ORANGE_HEX,
        after_pt=6,
    )
    doc.add_paragraph("")

    extracted: List[Dict[str, str]] = []
    if isinstance(extracted_rows, list) and extracted_rows:
        for x in extracted_rows:
            if not isinstance(x, dict):
                continue
            f = s(x.get("finding"))
            if not f:
                continue
            extracted.append(
                {
                    "finding": normalize_sentence(f),
                    "recommendation": normalize_sentence(x.get("recommendation")) or "—",
                }
            )

    if not extracted:
        p = doc.add_paragraph()
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        r = p.add_run("No major findings were captured to summarize.")
        set_run(r, FONT_BODY, 11, bold=False)
        return

    section0 = _safe_section(doc, 0)
    usable_emu = section_usable_width_emu(section0)
    usable_in = max(5.5, usable_emu / 914400.0)

    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    set_table_fixed_layout(table)
    set_table_width_from_section(table, doc, section_index=0)
    set_table_borders(table, color_hex=BORDER_HEX, size_eighths=BORDER_SZ_EIGHTHS)  # ✅ 1/4pt

    # ============================================================
    # Column widths (your requirement)
    #   No. = 0.40
    #   Severity = 0.69
    #   Finding == Recommendation (split remaining)
    # ============================================================
    w_no = 0.40
    w_sev = 0.69
    remaining = max(1.8, usable_in - (w_no + w_sev))
    w_find = remaining / 2.0
    w_rec = remaining / 2.0

    w_no, w_find, w_sev, w_rec = _scale_widths_to_fit(usable_in, [w_no, w_find, w_sev, w_rec])
    widths = [Inches(w_no), Inches(w_find), Inches(w_sev), Inches(w_rec)]

    headers = ["No.", "Finding", "Severity", "Recommendation / Corrective Action"]

    header_row = table.rows[0]
    set_row_cant_split(header_row, cant_split=True)
    set_row_height_exact(header_row, 0.25)

    for i, cell in enumerate(header_row.cells):
        try:
            cell.width = widths[i]
        except Exception:
            pass
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""

        shade_cell(cell, DARK_BLUE_HEX)
        set_cell_borders(cell, size_eighths=BORDER_SZ_EIGHTHS, color_hex=BORDER_HEX)  # ✅ 1/4pt
        set_cell_margins(cell, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        p = cell.paragraphs[0]
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)
        rr = p.add_run(headers[i])
        set_run(rr, FONT_BODY, FONT_SIZE_HEADER, True, HEADER_TEXT_COLOR)

    for idx, item in enumerate(extracted, start=1):
        finding = normalize_sentence(item.get("finding", ""))
        reco = normalize_sentence(item.get("recommendation", "")) or "—"

        row = table.add_row()
        set_row_cant_split(row, cant_split=False)
        set_row_height_at_least(row, 0.36)

        cells = row.cells
        for i, c in enumerate(cells):
            try:
                c.width = widths[i]
            except Exception:
                pass
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.text = ""
            set_cell_borders(c, size_eighths=BORDER_SZ_EIGHTHS, color_hex=BORDER_HEX)  # ✅ 1/4pt
            set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        p0 = cells[0].paragraphs[0]
        tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p0.add_run(str(idx)), FONT_BODY, FONT_SIZE_BODY, False)

        p1 = cells[1].paragraphs[0]
        tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p1.add_run(finding), FONT_BODY, FONT_SIZE_BODY, False)

        chosen = _resolve_severity(idx, finding, severity_by_no, severity_by_finding) or "Medium"
        p2 = cells[2].paragraphs[0]
        tight_paragraph(p2, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p2.add_run(chosen), FONT_BODY, FONT_SIZE_BODY, False)

        p3 = cells[3].paragraphs[0]
        tight_paragraph(p3, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p3.add_run(reco), FONT_BODY, FONT_SIZE_BODY, False)

    doc.add_paragraph("")
    doc.add_paragraph("")

    if not add_legend:
        return

    p_leg = doc.add_paragraph()
    tight_paragraph(p_leg, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
    set_run(p_leg.add_run("Legend:"), "Times New Roman", 11, True)

    legend_definitions = {
        "High": "Critical issue affecting functionality, safety, or compliance; requires immediate action.",
        "Medium": "Moderate issue affecting efficiency or performance; corrective action required.",
        "Low": "Minor issue with limited impact; corrective action recommended.",
    }

    present: List[str] = []
    found = set()
    for i in range(1, len(extracted) + 1):
        v = s(severity_by_no.get(i)).strip()
        if v:
            found.add(v.title())
    if not found:
        found = {"High", "Medium", "Low"}

    for sev in ["High", "Medium", "Low"]:
        if sev in found:
            present.append(sev)
    if not present:
        present = ["High", "Medium", "Low"]

    legend = doc.add_table(rows=1, cols=2)
    legend.autofit = False
    legend.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        legend.style = "Table Grid"
    except Exception:
        pass

    set_table_fixed_layout(legend)
    set_table_width_from_section(legend, doc, section_index=0)
    set_table_borders(legend, color_hex=BORDER_HEX, size_eighths=BORDER_SZ_EIGHTHS)  # ✅ 1/4pt

    r0 = legend.rows[0]
    set_row_cant_split(r0, cant_split=True)
    set_row_height_exact(r0, 0.25)

    h0, h1c = r0.cells
    lw0 = 1.10
    lw1 = max(2.5, usable_in - lw0)
    lw0, lw1 = _scale_widths_to_fit(usable_in, [lw0, lw1])

    h0.width = Inches(lw0)
    h1c.width = Inches(lw1)

    for c, txt in ((h0, "Severity"), (h1c, "Definition")):
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c, DARK_BLUE_HEX)
        set_cell_borders(c, size_eighths=BORDER_SZ_EIGHTHS, color_hex=BORDER_HEX)  # ✅ 1/4pt
        set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        pp = c.paragraphs[0]
        tight_paragraph(pp, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)
        set_run(pp.add_run(txt), FONT_BODY, 10, True, HEADER_TEXT_COLOR)

    for sev in present:
        rr = legend.add_row()
        set_row_cant_split(rr, cant_split=False)
        set_row_height_at_least(rr, 0.25)

        c0, c1 = rr.cells
        c0.width = Inches(lw0)
        c1.width = Inches(lw1)

        for c in (c0, c1):
            c.text = ""
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(c, size_eighths=BORDER_SZ_EIGHTHS, color_hex=BORDER_HEX)  # ✅ 1/4pt
            set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        p0 = c0.paragraphs[0]
        tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p0.add_run(sev), FONT_BODY, 10, False)

        p1 = c1.paragraphs[0]
        tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p1.add_run(legend_definitions.get(sev, "")), FONT_BODY, 10, False)
