# src/report_sections/toc_page.py
from __future__ import annotations

from typing import Optional, Union

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.text.paragraph import Paragraph


# ============================================================
# Inline helpers (replacing src.report_sections._word_common)
# ============================================================
def set_run(
    run,
    font: str,
    size: Union[int, float],
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    """Safe run styling with eastAsia font set."""
    try:
        run.font.name = font
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass
    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass
    run.bold = bool(bold)
    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: Union[int, float] = 0,
    after_pt: Union[int, float] = 0,
    line_spacing: float = 1.0,
) -> None:
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def _set_paragraph_bottom_border(
    paragraph: Paragraph,
    *,
    color_hex: str,
    size_eighths: int = 12,  # 12 => 1.5pt
    space: int = 2,
) -> None:
    """Adds/updates a bottom border line on a paragraph (idempotent)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), str(color_hex).replace("#", ""))


def title_with_orange_line(
    doc: Document,
    *,
    text: str,
    font: str = "Cambria",
    size: int = 16,
    color: RGBColor = RGBColor(0, 112, 192),
    orange_hex: str = "ED7D31",
    after_pt: Union[int, float] = 10,
) -> Paragraph:
    """
    Title paragraph with orange underline — NOT a heading (so it won't appear in TOC).
    """
    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1.0)

    r = p.add_run(str(text))
    set_run(r, font, size, bold=True, color=color)

    _set_paragraph_bottom_border(p, color_hex=orange_hex, size_eighths=12, space=2)
    return p


def add_toc_field(
    doc: Document,
    *,
    levels: str = "1-3",
    hyperlinks: bool = True,
    hide_page_numbers_in_web_layout: bool = False,
) -> None:
    """
    Inserts a Word TOC field code.
    - levels: like "1-3"
    - hyperlinks: adds \\h
    - hide_page_numbers_in_web_layout: adds \\z (Word web layout switch)
    """
    # Field instructions
    instr = f'TOC \\o "{levels}"'
    if hyperlinks:
        instr += " \\h"
    if hide_page_numbers_in_web_layout:
        instr += " \\z"
    instr += " \\u"  # use outline levels (safe default)

    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1.0)

    r = p.add_run()

    # <w:fldChar w:fldCharType="begin"/>
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fld_begin)

    # <w:instrText xml:space="preserve">...</w:instrText>
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    r._r.append(instr_text)

    # <w:fldChar w:fldCharType="separate"/>
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r._r.append(fld_sep)

    # Placeholder text (Word updates it when fields are updated)
    r2 = p.add_run("Table of Contents will be generated here.")
    set_run(r2, "Times New Roman", 11, bold=False)

    # <w:fldChar w:fldCharType="end"/>
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p.runs[-1]._r.append(fld_end)


# ============================================================
# Section constants
# ============================================================
TITLE_TEXT = "Table of Contents"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16
TITLE_BLUE = RGBColor(0, 112, 192)

BODY_FONT = "Times New Roman"
BODY_SIZE = 11


def add_toc_page(
    doc: Document,
    *,
    toc_levels: str = "1-3",
    include_hyperlinks: bool = True,
    hide_page_numbers_in_web_layout: bool = False,
) -> None:
    """
    Adds a TOC page using a Word field code.
    NOTE: This page title is NOT a Heading, so it won't appear in TOC.
    """
    doc.add_page_break()

    title_with_orange_line(
        doc,
        text=TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=TITLE_BLUE,
        orange_hex="ED7D31",
        after_pt=10,
    )

    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=8, line_spacing=1)
    set_run(
        p.add_run("Right-click the table and select “Update Field” in Word to refresh page numbers."),
        BODY_FONT,
        BODY_SIZE,
        bold=False,
    )

    add_toc_field(
        doc,
        levels=toc_levels,
        hyperlinks=include_hyperlinks,
        hide_page_numbers_in_web_layout=hide_page_numbers_in_web_layout,
    )
