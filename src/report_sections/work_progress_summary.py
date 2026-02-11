from __future__ import annotations

import re
from typing import List, Optional, Union, Any, Dict

from docx.document import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


# ============================================================
# Inline helpers (clean + stable)
# ============================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: Union[int, float] = 0,
    after_pt: Union[int, float] = 0,
    line_spacing: float = 1.0,
) -> None:
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def set_run(
    run,
    font: str,
    size: Union[int, float],
    *,
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    # Font name
    try:
        run.font.name = font
    except Exception:
        pass

    # EastAsia font (important for some Word envs)
    try:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font)
    except Exception:
        pass

    # Size
    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass

    run.bold = bool(bold)

    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


def set_paragraph_bottom_border(
    paragraph: Paragraph,
    *,
    color_hex: str,
    size_eighths: int = 12,  # 12 => 1.5pt
    space: int = 2,
) -> None:
    """
    Bottom border under a paragraph (orange underline).
    size_eighths is in 1/8 pt units (Word).
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), s(color_hex).replace("#", "") or "000000")


def add_heading(
    doc: Document,
    text: str,
    *,
    level: int = 1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font: str = "Cambria",
    size: int = 16,
    bold: bool = True,
    color: Optional[RGBColor] = None,
) -> Paragraph:
    """
    Adds a real Word heading (TOC-safe). Keeps style intact.
    """
    lvl = max(1, min(int(level), 9))
    p = doc.add_paragraph()
    try:
        p.style = f"Heading {lvl}"
    except Exception:
        pass

    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1.0)
    r = p.add_run(s(text))
    set_run(r, font, size, bold=bold, color=color)
    return p


def add_section_title_h1(
    doc: Document,
    text: str,
    *,
    font: str = "Cambria",
    size: int = 16,
    # IMPORTANT: keep Heading 1 "natural" color by default => None
    color: Optional[RGBColor] = None,
    orange_hex: str = "ED7D31",
    after_pt: Union[int, float] = 6,
) -> Paragraph:
    """
    Standard H1 title:
      - Heading 1 (TOC-safe)
      - size 16
      - orange underline on same paragraph
      - if color=None => keep Word style default color
    """
    p = add_heading(
        doc,
        s(text),
        level=1,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font=font,
        size=int(size),
        bold=True,
        color=color,
    )
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1.0)
    set_paragraph_bottom_border(p, color_hex=orange_hex, size_eighths=12, space=2)
    return p


# -------------------------
# Table helpers
# -------------------------
def set_table_fixed_layout(table) -> None:
    """
    Prevent Word from auto-resizing columns.
    """
    try:
        table.allow_autofit = False  # type: ignore[attr-defined]
    except Exception:
        pass
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def set_table_borders(table, *, color_hex: str = "A6A6A6", size_eighths: int = 2) -> None:
    """
    Borders for the whole table.

    ✅ 1/4 pt borders:
      Word border size uses 1/8 pt units:
      0.25pt * 8 = 2  -> w:sz="2"
    """
    color_hex = s(color_hex).replace("#", "") or "A6A6A6"

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(int(size_eighths)))  # ✅ 1/4pt = 2
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def shade_cell(cell, fill_hex: str) -> None:
    fill_hex = s(fill_hex).replace("#", "") or "FFFFFF"
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_margins(cell, *, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120) -> None:
    """
    Set cell internal margins. DXA (twips): 1440 twips = 1 inch.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _set(side: str, val: int) -> None:
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")

    _set("top", int(top_dxa))
    _set("bottom", int(bottom_dxa))
    _set("start", int(left_dxa))
    _set("end", int(right_dxa))


def set_row_cant_split(row, *, cant_split: bool = True) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    existing = trPr.find(qn("w:cantSplit"))
    if cant_split:
        if existing is None:
            trPr.append(OxmlElement("w:cantSplit"))
    else:
        if existing is not None:
            trPr.remove(existing)


def write_cell_text(
    cell,
    text: Any,
    *,
    font: str = "Times New Roman",
    size: int = 11,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    valign: Optional[int] = None,
) -> None:
    """
    Write single-paragraph text with alignment + optional vertical alignment.
    """
    cell.text = ""
    if valign is not None:
        try:
            cell.vertical_alignment = valign
        except Exception:
            pass

    p = cell.paragraphs[0]
    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1.0)
    r = p.add_run(s(text))
    set_run(r, font, size, bold=bold, color=None)


def strip_heading_numbering(text: str) -> str:
    """
    Remove leading numbering like:
      '1.2 Title', '5) Title', '3 - Title', '2.    Title'
    """
    t = s(text)
    if not t:
        return ""
    t = re.sub(r"^\s*\(?\d+(\.\d+)*\)?\s*[\.\)\-:]\s*", "", t)
    t = re.sub(r"^\s*\d+(\.\d+)*\s+", "", t)
    return t.strip()


# -----------------
# Style (match report rule)
# -----------------
TITLE_FONT = "Cambria"
TITLE_SIZE = 16
ORANGE_HEX = "ED7D31"

BODY_FONT = "Times New Roman"
BODY_SIZE = 11

HEADER_FILL_HEX = "D9E2F3"
BORDER_HEX = "A6A6A6"  # border color
BORDER_SZ_EIGHTHS = 2  # ✅ 1/4 pt


# ============================================================
# COLUMN WIDTHS (INCHES)
# ============================================================
# NOTE: use float inches here for clarity; apply via Inches() below
COL_W_NO_IN = 0.40
COL_W_ACT_IN = 3.63
COL_W_PLAN_IN = 0.81
COL_W_ACH_IN = 0.81
COL_W_PROG_IN = 0.81
COL_W_REM_IN = 0.81

COLUMN_WIDTHS_IN = [COL_W_NO_IN, COL_W_ACT_IN, COL_W_PLAN_IN, COL_W_ACH_IN, COL_W_PROG_IN, COL_W_REM_IN]
# ============================================================


def _apply_column_widths(table, widths_in: List[float]) -> None:
    """
    Apply widths to table columns and all cells.
    Uses fixed layout => Word won't resize.
    """
    for i, w_in in enumerate(widths_in):
        w = Inches(float(w_in))
        table.columns[i].width = w
    for row in table.rows:
        for i, w_in in enumerate(widths_in):
            row.cells[i].width = Inches(float(w_in))


def _normalize_progress_for_doc(progress_text: Any) -> str:
    """
    Keep user's value as much as possible.
    - If contains '%' -> keep
    - If plain number -> clamp 0..100 and add '%'
    - Else keep raw text
    """
    t = s(progress_text)
    if not t:
        return ""
    if "%" in t:
        return t

    tt = t.replace(" ", "")
    if re.fullmatch(r"\d+(\.\d+)?", tt):
        try:
            n = int(float(tt))
            n = max(0, min(100, n))
            return f"{n}%"
        except Exception:
            return t

    return t


def _normalize_rows(rows: Optional[List[Dict[str, Any]]]) -> List[Dict[str, str]]:
    """
    Accepts rows like:
      [{"Activities": "...", "Planned": "...", "Achieved": "...", "Progress": "...", "Remarks": "..."}]
    Returns cleaned list (strings).
    """
    out: List[Dict[str, str]] = []
    for r in (rows or []):
        if not isinstance(r, dict):
            continue

        act = s(r.get("Activities"))
        planned = s(r.get("Planned"))
        achieved = s(r.get("Achieved"))
        progress = _normalize_progress_for_doc(r.get("Progress"))
        remarks = s(r.get("Remarks"))

        # skip fully empty rows
        if not act and not any([planned, achieved, progress, remarks]):
            continue

        out.append(
            {
                "Activities": act,
                "Planned": planned,
                "Achieved": achieved,
                "Progress": progress,
                "Remarks": remarks,
            }
        )
    return out


def add_work_progress_summary_during_visit(
    doc: Document,
    *,
    activity_titles_from_section5: List[str],
    title_text: str = "4.    Work Progress Summary during the Visit.",
    rows: Optional[List[Dict[str, Any]]] = None,
) -> None:
    """
    Section 4 table.

    ✅ If `rows` provided (from Step 5), fills the table.
    ✅ Else uses titles list and leaves numeric columns blank.

    ✅ Formatting:
      - Heading 1 keeps its native (style) color
      - table fixed layout
      - borders = 1/4 pt (w:sz=2)
      - Header centered with fill
      - Body: No/Planned/Achieved/Progress centered
    """

    # Keep Heading 1 natural color (color=None)
    add_section_title_h1(
        doc,
        s(title_text),
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=None,
        orange_hex=ORANGE_HEX,
        after_pt=6,
    )

    # spacing after title
    psp = doc.add_paragraph("")
    tight_paragraph(psp, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=6, line_spacing=1.0)

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    set_table_fixed_layout(table)
    set_table_borders(table, color_hex=BORDER_HEX, size_eighths=BORDER_SZ_EIGHTHS)

    _apply_column_widths(table, COLUMN_WIDTHS_IN)

    headers = ["No.", "Activities", "Planned", "Achieved", "Progress", "Remarks"]

    hdr = table.rows[0]
    set_row_cant_split(hdr, cant_split=True)

    for i, cell in enumerate(hdr.cells):
        try:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            pass

        shade_cell(cell, HEADER_FILL_HEX)
        set_cell_margins(cell, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        write_cell_text(
            cell,
            headers[i],
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            valign=WD_ALIGN_VERTICAL.CENTER,
        )

    # ---- Body source
    data_rows = _normalize_rows(rows)

    if not data_rows:
        clean_acts = [
            strip_heading_numbering(s(t)).strip()
            for t in (activity_titles_from_section5 or [])
            if s(t).strip()
        ]
        if not clean_acts:
            clean_acts = ["", "", ""]
        data_rows = [{"Activities": act, "Planned": "", "Achieved": "", "Progress": "", "Remarks": ""} for act in clean_acts]

    # ---- Body rows
    for idx, rr in enumerate(data_rows, start=1):
        row = table.add_row()
        set_row_cant_split(row, cant_split=False)

        _apply_column_widths(table, COLUMN_WIDTHS_IN)

        cells = row.cells
        for c in cells:
            try:
                c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            except Exception:
                pass
            set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        # No.
        write_cell_text(
            cells[0],
            str(idx),
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            valign=WD_ALIGN_VERTICAL.CENTER,
        )

        # Activities
        write_cell_text(
            cells[1],
            rr.get("Activities", ""),
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            valign=WD_ALIGN_VERTICAL.TOP,
        )

        # Planned
        write_cell_text(
            cells[2],
            rr.get("Planned", ""),
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            valign=WD_ALIGN_VERTICAL.CENTER,
        )

        # Achieved
        write_cell_text(
            cells[3],
            rr.get("Achieved", ""),
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            valign=WD_ALIGN_VERTICAL.CENTER,
        )

        # Progress
        write_cell_text(
            cells[4],
            rr.get("Progress", ""),
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            valign=WD_ALIGN_VERTICAL.CENTER,
        )

        # Remarks
        write_cell_text(
            cells[5],
            rr.get("Remarks", ""),
            font=BODY_FONT,
            size=BODY_SIZE,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            valign=WD_ALIGN_VERTICAL.TOP,
        )

    # final safety width apply
    _apply_column_widths(table, COLUMN_WIDTHS_IN)
